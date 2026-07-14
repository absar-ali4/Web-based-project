import io
import os
import re
import time
import random
import base64
import threading
import urllib.parse
from contextlib import contextmanager
from pathlib import Path

import cv2
import numpy as np
import pandas as pd
import requests
import streamlit as st
import streamlit.components.v1 as components
from PIL import Image, ImageDraw, ImageEnhance, ImageFilter, ImageFont
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ------------------------------------------------------------------
# PAGE CONFIG
# ------------------------------------------------------------------
def _build_favicon(size: int = 64) -> Image.Image:
    """Renders the same neural-node brand mark used in the navbar and hero
    heading (see `brand_logo_svg`) as a small transparent PNG, so the browser
    tab icon matches the in-app logo instead of a generic emoji."""
    scale = size / 32
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    def pt(x, y):
        return (x * scale, y * scale)

    # Soft glow disc behind the mark, mirroring the SVG's radial backdrop.
    draw.ellipse([pt(1, 1), pt(31, 31)], fill=(56, 189, 248, 40))

    # Connecting neural links (same coordinates as brand_logo_svg).
    line_color = (56, 189, 248, 235)
    lw = max(1, round(1.7 * scale))
    draw.line([pt(16, 9.6), pt(16, 13.6)], fill=line_color, width=lw)
    draw.line([pt(14.1, 17.6), pt(9.4, 20.4)], fill=line_color, width=lw)
    draw.line([pt(17.9, 17.6), pt(262.6, 20.4)], fill=line_color, width=lw)

    # Gradient-style nodes (cyan at top fading to violet at the base).
    for (cx, cy), r, color in [
        ((16, 7.2), 2.3, (125, 211, 252, 255)),
        ((7.6, 22.4), 2.3, (124, 58, 237, 255)),
        ((24.4, 22.4), 2.3, (124, 58, 237, 255)),
        ((16, 16), 2.5, (56, 189, 248, 255)),
    ]:
        draw.ellipse([pt(cx - r, cy - r), pt(cx + r, cy + r)], fill=color)
    return img


st.set_page_config(page_title="NeuralCraft", page_icon=_build_favicon(), layout="wide")

# ------------------------------------------------------------------
# THEME BOOTSTRAP (must exist before the stylesheet is built)
# ------------------------------------------------------------------
st.session_state.setdefault("theme", "dark")

PALETTES = {
    "dark": dict(
        app_bg="linear-gradient(160deg, #050710 0%, #0a0e1a 50%, #000000 100%)",
        navbar_bg="rgba(8,10,20,.82)", navbar_border="rgba(56,189,248,.25)",
        panel_bg="rgba(10,13,24,.6)", panel_border="rgba(255,255,255,.08)",
        panel_border_hover="rgba(56,189,248,.4)",
        text_primary="#f8fafc", text_secondary="#a3adc2", text_inverse="#05070d",
        hero_bg="linear-gradient(135deg, rgba(8,10,20,.97), rgba(16,20,36,.92))",
        hero_title="linear-gradient(135deg, #7dd3fc 0%, #e0e7ff 45%, #c4b5fd 100%)",
        radio_bg="rgba(8,10,20,.7)", card_bg="linear-gradient(135deg, rgba(16,20,36,.8), rgba(8,10,20,.92))",
        toggle_icon="☀️",
    ),
    "light": dict(
        app_bg="linear-gradient(160deg, #ffffff 0%, #f4f6fb 55%, #ffffff 100%)",
        navbar_bg="rgba(255,255,255,.92)", navbar_border="rgba(37,99,235,.16)",
        panel_bg="rgba(255,255,255,.85)", panel_border="rgba(15,23,42,.07)",
        panel_border_hover="rgba(37,99,235,.32)",
        text_primary="#0b1220", text_secondary="#52607a", text_inverse="#ffffff",
        hero_bg="linear-gradient(135deg, rgba(255,255,255,.98), rgba(244,246,251,.95))",
        hero_title="linear-gradient(135deg, #1d4ed8 0%, #4338ca 55%, #6d28d9 100%)",
        radio_bg="rgba(255,255,255,.88)", card_bg="linear-gradient(135deg, rgba(255,255,255,.94), rgba(244,246,251,.94))",
        toggle_icon="🌙",
    ),
}
P = PALETTES[st.session_state.theme]

# ------------------------------------------------------------------
# STYLES
# ------------------------------------------------------------------
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@500;700&family=Plus+Jakarta+Sans:wght@400;500;600;700&family=Georgia&display=swap');

html, body, .stApp {{ font-family: 'Plus Jakarta Sans', sans-serif; }}

:root {{
    --space-lg: clamp(18px, 3vw, 32px);
    --space-md: clamp(12px, 2vw, 22px);
    --radius: 18px;
    --accent: #2563eb;
    --accent-2: #38bdf8;
    --nav-font: Georgia, 'Times New Roman', 'Bell MT', serif;
    --heading-font: 'Bell MT', Georgia, 'Times New Roman', serif;
    --text-primary: {P['text_primary']};
    --text-secondary: {P['text_secondary']};
    --ease: cubic-bezier(.4, 0, .2, 1);
    --dur-fast: .2s;
    --dur-med: .35s;
    --dur-slow: .5s;
}}

* {{ scroll-behavior: smooth; box-sizing: border-box; -webkit-tap-highlight-color: transparent; }}
html, body {{ overflow-x: hidden; }}
img, svg {{ max-width: 100%; height: auto; }}

/* Reduced-motion users get instant state changes across the board. Individual
   components each declare their own scoped transition (see below) rather than
   a blanket "transition on every element" rule, which is expensive to
   recompute on Streamlit's frequent full-tree reruns and was the main source
   of jank/flicker/stutter. */
@media (prefers-reduced-motion: reduce) {{
    * {{ transition: none !important; animation: none !important; }}
}}

/* Touch/coarse-pointer devices have no real ":hover" -- the browser fakes it
   on tap, which leaves decorative lift/scale effects visually "stuck" until
   the user taps elsewhere. Neutralizing just the transform on those devices
   keeps every hover effect fully intact for mouse/trackpad users while
   making tap interactions feel clean and responsive on phones/tablets. */
@media (hover: none) {{
    .st-key-navbar_row:hover, .hero-box:hover, .profile-card:hover, .profile-photo:hover,
    .feature-card:hover, .feature-card:hover .feature-icon, .nav-brand:hover .nav-logo,
    .center-col img:hover, [data-testid="stVerticalBlockBorderWrapper"]:hover {{
        transform: none !important;
    }}
}}

.stApp {{ background: {P['app_bg']} !important; transition: background var(--dur-slow) var(--ease); }}
.stApp, .stApp p, .stApp span, .stApp label {{ color: var(--text-primary); transition: color var(--dur-med) var(--ease); }}

/* ---------- Navbar row (brand + navigation + appearance menu) ---------- */
.st-key-navbar_row {{
    background: {P['navbar_bg']};
    border: 1px solid {P['navbar_border']};
    border-radius: 16px;
    padding: 10px 22px;
    margin-bottom: 24px;
    backdrop-filter: blur(14px);
    transition: background var(--dur-slow) var(--ease), border-color var(--dur-med) var(--ease), box-shadow var(--dur-med) var(--ease);
}}
.st-key-navbar_row:hover {{ border-color: rgba(56,189,248,.4); box-shadow: 0 10px 30px rgba(0,0,0,.18); }}
.st-key-navbar_row [data-testid="stHorizontalBlock"] {{
    align-items: center !important;
    flex-wrap: nowrap !important;
}}
/* Force a single row (brand | theme toggle | nav) at every breakpoint, incl. mobile */
.st-key-navbar_row [data-testid="stHorizontalBlock"] > [data-testid="column"] {{
    min-width: 0 !important;
}}
.st-key-navbar_row [data-testid="stHorizontalBlock"] > [data-testid="column"]:nth-of-type(2),
.st-key-navbar_row [data-testid="stHorizontalBlock"] > [data-testid="column"]:nth-of-type(3) {{
    width: auto !important;
    flex: 0 0 auto !important;
}}

.nav-brand {{ display: flex; align-items: center; gap: 10px; height: 48px; min-width: 0; overflow: hidden; perspective: 500px; }}
.nav-logo {{
    display: inline-flex; align-items: center; justify-content: center;
    line-height: 1; flex-shrink: 0; transform-style: preserve-3d;
    filter: drop-shadow(0 0 6px rgba(56,189,248,.35));
    animation: nc-logo-spin-3d 7s linear infinite;
    transition: filter .3s ease, animation-duration .3s ease;
}}
.nav-logo svg {{ display: block; }}
@keyframes nc-logo-spin-3d {{
    from {{ transform: rotateY(0deg); }}
    to   {{ transform: rotateY(360deg); }}
}}
.nav-brand:hover .nav-logo {{ animation-duration: 1.6s; filter: drop-shadow(0 0 10px rgba(56,189,248,.55)); }}
@media (prefers-reduced-motion: reduce) {{
    .nav-logo {{ animation: none; }}
}}
.nav-title {{
    font-family: var(--heading-font);
    font-size: clamp(15px, 2.4vw, 24px);
    font-weight: 700;
    color: var(--text-primary);
    letter-spacing: .2px;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
}}
/* Belt-and-braces: never let the navbar row itself wrap horizontally, so the
   hamburger/⋮ pair can never slip below the logo on narrow devices. Only
   the X axis is clipped -- Y stays visible so the Appearance/hamburger
   popovers (which open downward) are never cut off on mobile. */
.st-key-navbar_row {{ overflow-x: hidden; overflow-y: visible; }}
.st-key-navbar_row [data-testid="stHorizontalBlock"] {{ overflow: visible; }}

/* Desktop nav buttons -- elegant serif labels, plain transparent pill-hover links */
.st-key-nav_desktop {{ display: flex; justify-content: flex-end; }}
.st-key-nav_desktop [data-testid="stHorizontalBlock"] {{ gap: 4px !important; justify-content: flex-end; flex-wrap: nowrap !important; }}
.st-key-nav_desktop .stButton > button {{
    background: transparent !important;
    border: 1px solid transparent !important;
    color: var(--text-secondary) !important;
    font-family: var(--nav-font) !important;
    font-weight: 700;
    font-size: clamp(12.5px, 1.3vw, 15.5px);
    white-space: nowrap;
    box-shadow: none !important;
    padding: 8px clamp(10px, 1.6vw, 18px);
    transition: background .25s ease, color .25s ease, transform .2s ease;
}}
.st-key-nav_desktop .stButton > button:hover {{
    background: rgba(56,189,248,.12) !important;
    color: var(--text-primary) !important;
    transform: translateY(-1px);
}}
.st-key-nav_desktop .stButton > button[kind="primary"] {{
    background: linear-gradient(135deg, var(--accent), #1d4ed8) !important;
    color: #fff !important;
    box-shadow: 0 4px 14px rgba(37,99,235,.4) !important;
}}
/* Serif nav font also applied to the mobile popover items */
div[data-testid="stPopoverBody"] .stButton > button {{ font-family: var(--nav-font) !important; font-weight: 700; }}

/* Mobile hamburger trigger -- pinned top-right of the navbar */
.st-key-nav_mobile {{ display: flex; justify-content: flex-end; }}
.st-key-nav_mobile button {{
    background: rgba(56,189,248,.1) !important;
    border: 1px solid rgba(56,189,248,.3) !important;
    color: var(--text-primary) !important;
    font-size: 18px !important;
    border-radius: 10px !important;
    padding: 6px 14px !important;
    transition: transform var(--dur-fast) var(--ease), background var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease);
}}
.st-key-nav_mobile button:hover {{
    transform: scale(1.06);
    background: rgba(56,189,248,.2) !important;
    box-shadow: 0 6px 16px rgba(56,189,248,.25);
}}
.st-key-nav_mobile button:active {{ transform: scale(.96); }}

/* Appearance (⋮) menu trigger -- visible at every breakpoint, holds Light/Dark options */
.st-key-theme_menu_wrap {{ display: flex; justify-content: flex-end; }}
.st-key-theme_menu_wrap button {{
    background: rgba(56,189,248,.1) !important;
    border: 1px solid rgba(56,189,248,.3) !important;
    color: var(--text-primary) !important;
    font-size: 17px !important;
    font-weight: 700 !important;
    border-radius: 10px !important;
    padding: 6px 12px !important;
    line-height: 1 !important;
    transition: transform var(--dur-fast) var(--ease), background var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease);
}}
.st-key-theme_menu_wrap button:hover {{
    transform: scale(1.08);
    background: rgba(56,189,248,.2) !important;
    box-shadow: 0 6px 16px rgba(56,189,248,.25);
}}
.st-key-theme_menu_wrap button:active {{ transform: scale(.94); }}
div[data-testid="stPopoverBody"] {{
    animation: menu-in .18s var(--ease);
}}
@keyframes menu-in {{
    from {{ opacity: 0; transform: translateY(-4px); }}
    to {{ opacity: 1; transform: translateY(0); }}
}}

/* Responsive visibility: swap desktop nav <-> mobile hamburger */
.mobile-appearance-block {{ display: none; }}
@media (max-width: 640px) {{
    .st-key-nav_desktop {{ display: none !important; }}
    .desktop-only {{ display: none !important; }}
    .st-key-navbar_row {{ padding: 8px 14px; }}
    .nav-title {{ font-size: 17px; }}
    /* Appearance moves inside the ☰ hamburger menu on mobile/small devices */
    .st-key-theme_menu_wrap {{ display: none !important; }}
    .mobile-appearance-block {{ display: block; border-top: 1px dashed {P['panel_border']}; margin-top: 4px; }}
}}
@media (min-width: 641px) {{
    .st-key-nav_mobile {{ display: none !important; }}
}}

/* ---------- Hero ---------- */
.hero-box {{
    background: {P['hero_bg']};
    border: 1px solid rgba(56,189,248,.2);
    border-radius: 24px;
    padding: var(--space-lg);
    text-align: center;
    box-shadow: 0 20px 40px rgba(0,0,0,.25), 0 0 50px rgba(37,99,235,.12);
    margin-bottom: 28px;
    transition: background var(--dur-slow) var(--ease), box-shadow var(--dur-med) var(--ease), transform var(--dur-med) var(--ease);
    animation: fade-in-up .5s var(--ease);
}}
.hero-box:hover {{ transform: translateY(-3px); box-shadow: 0 25px 50px rgba(0,0,0,.35), 0 0 60px rgba(56,189,248,.2); }}
.hero-title {{
    font-family: var(--heading-font); font-size: clamp(22px, 3vw, 32px); font-weight: 700;
    margin: 0 0 8px 0; background: {P['hero_title']};
    -webkit-background-clip: text; background-clip: text; color: transparent;
    display: flex; align-items: center; justify-content: center; gap: 10px; flex-wrap: wrap;
}}
/* ---------- Hero heading typewriter effect ---------- */
/* Types itself in on page load. A blinking cursor sits right after the text
   while it's typing, then stops blinking and fades away once the animation
   completes. */
.typewriter {{
    display: inline-block;
    overflow: hidden;
    white-space: nowrap;
    vertical-align: bottom;
    max-width: 100%;
    width: 0;
    animation: nc-typing 1.3s steps(37, end) .1s forwards;
}}
@keyframes nc-typing {{ from {{ width: 0; }} to {{ width: 37ch; }} }}
.typewriter-cursor {{
    display: inline-block;
    width: 3px;
    height: 0.85em;
    margin-left: 3px;
    vertical-align: -0.1em;
    background: var(--accent-2);
    border-radius: 1px;
    animation: nc-caret-blink .6s steps(1, end) 3 .1s,
               nc-caret-hide 0s linear 1.5s forwards;
}}
@keyframes nc-caret-blink {{ 50% {{ opacity: 0; }} }}
@keyframes nc-caret-hide {{ to {{ opacity: 0; }} }}
@media (prefers-reduced-motion: reduce) {{
    .typewriter {{ width: 37ch; animation: none; }}
    .typewriter-cursor {{ display: none; }}
}}
.hero-subtitle {{ font-size: clamp(13px, 1.5vw, 15px); color: var(--text-secondary); margin: 0; }}
@media (max-width: 640px) {{ .hero-box {{ display: none; }} }}

/* ---------- Segmented controls (mode switches) ---------- */
.mode-switch-wrap {{ display: flex; justify-content: center; margin-bottom: 32px; }}
div[data-testid="stRadio"] > div[role="radiogroup"] {{
    display: flex; flex-wrap: wrap; justify-content: center; gap: 6px;
    background: {P['radio_bg']}; border: 1px solid rgba(56,189,248,.25);
    border-radius: 16px; padding: 6px; transition: background .4s ease;
}}
div[data-testid="stRadio"] input[type="radio"] {{ position: absolute; opacity: 0; width: 0; height: 0; }}
div[data-testid="stRadio"] label {{
    margin: 0 !important; border-radius: 12px; padding: 10px 22px; cursor: pointer;
    color: var(--text-secondary); font-family: var(--heading-font); font-weight: 700;
    transition: background .25s ease, color .25s ease, transform .15s ease;
}}
div[data-testid="stRadio"] label:hover {{ color: var(--text-primary); transform: translateY(-1px); }}
div[data-testid="stRadio"] label:has(input:checked) {{
    background: linear-gradient(135deg, var(--accent), #1d4ed8);
    color: #fff !important; box-shadow: 0 4px 14px rgba(37,99,235,.4);
}}

/* ---------- Panel cards ---------- */
[data-testid="stVerticalBlockBorderWrapper"] {{
    border-radius: var(--radius) !important;
    border: 1px solid {P['panel_border']} !important;
    background: {P['panel_bg']} !important;
    padding: var(--space-md) !important;
    transition: background var(--dur-slow) var(--ease), border-color var(--dur-med) var(--ease),
                transform var(--dur-med) var(--ease), box-shadow var(--dur-med) var(--ease) !important;
    animation: fade-in-up .45s var(--ease);
}}
[data-testid="stVerticalBlockBorderWrapper"]:hover {{
    border-color: {P['panel_border_hover']} !important;
    box-shadow: 0 14px 32px rgba(0,0,0,.16);
}}

.panel-header {{
    display: flex; align-items: center; gap: 10px;
    font-family: var(--heading-font); font-size: clamp(17px, 2vw, 20px); font-weight: 700;
    color: var(--accent-2); border-left: 4px solid var(--accent); padding-left: 12px; margin-bottom: 16px;
    transition: border-color var(--dur-med) var(--ease);
}}
.panel-header-icon {{ display: inline-flex; flex-shrink: 0; }}
.panel-header-icon svg {{ width: 20px; height: 20px; stroke: var(--accent-2); }}

/* ---------- Buttons ---------- */
.stButton > button {{
    border-radius: 12px; font-weight: 600; padding: 10px 24px;
    background: linear-gradient(135deg, var(--accent), #1d4ed8); color: #fff;
    border: 1px solid rgba(255,255,255,.15);
    transition: transform var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease),
                filter var(--dur-fast) var(--ease), background var(--dur-med) var(--ease);
    will-change: transform;
}}
.stButton > button:hover {{ transform: translateY(-2px) scale(1.015); box-shadow: 0 10px 24px rgba(37,99,235,.42); filter: brightness(1.06); }}
.stButton > button:active {{ transform: translateY(0) scale(.98); box-shadow: 0 4px 10px rgba(37,99,235,.3); }}

/* Secondary / outline buttons get their own subtle lift */
.stButton > button[kind="secondary"] {{ transition: transform var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease), background var(--dur-fast) var(--ease); }}
.stButton > button[kind="secondary"]:hover {{ transform: translateY(-1px); }}

.stDownloadButton > button {{
    transition: transform var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease), filter var(--dur-fast) var(--ease);
}}
.stDownloadButton > button:hover {{ transform: translateY(-2px) scale(1.015); box-shadow: 0 10px 22px rgba(37,99,235,.35); filter: brightness(1.06); }}
.stDownloadButton > button:active {{ transform: translateY(0) scale(.98); }}

/* ---------- Centered image wrapper ---------- */
.center-col {{ display: flex; flex-direction: column; align-items: center; text-align: center; width: 100%; }}
.center-col img {{ border-radius: 14px; transition: transform var(--dur-med) var(--ease), box-shadow var(--dur-med) var(--ease); animation: fade-in .5s var(--ease); }}
.center-col img:hover {{ transform: scale(1.02); box-shadow: 0 16px 34px rgba(0,0,0,.22); }}
@keyframes fade-in {{ from {{ opacity: 0; transform: translateY(6px); }} to {{ opacity: 1; transform: translateY(0); }} }}
@keyframes fade-in-up {{ from {{ opacity: 0; transform: translateY(10px); }} to {{ opacity: 1; transform: translateY(0); }} }}

/* Stacked module cards ease in with a subtle stagger for a premium feel */
.st-key-studio_columns [data-testid="stVerticalBlockBorderWrapper"]:nth-of-type(1) {{ animation-delay: 0s; }}
.st-key-studio_columns [data-testid="stVerticalBlockBorderWrapper"]:nth-of-type(2) {{ animation-delay: .08s; }}

/* Accessible, theme-aware focus rings for keyboard navigation */
.stButton > button:focus-visible, .stDownloadButton > button:focus-visible,
.stTextArea textarea:focus-visible, .stTextInput input:focus-visible {{
    outline: 2px solid var(--accent-2) !important; outline-offset: 2px;
}}

/* ---------- Sketch preview panels (balanced, capped size) ---------- */
.preview-box {{ max-width: 460px; margin: 0 auto; }}
.preview-box .center-col img {{ max-width: 100%; }}
.preview-box--sketch {{ max-width: 100%; }}
.density-label {{
    font-family: var(--heading-font); font-weight: 700; font-size: 14px;
    color: var(--accent-2); margin-bottom: 4px;
}}
.density-panel {{
    display: flex; flex-direction: column; justify-content: center; height: 100%;
    padding-right: 6px;
}}

/* ---------- Shading control panel (sits immediately left of the sketch preview) ---------- */
.shading-panel {{
    display: flex; flex-direction: column; justify-content: center; height: 100%;
    padding-right: 6px;
    border-right: 1px dashed {P['panel_border']};
    transition: border-color var(--dur-med) var(--ease);
}}
.shading-label {{
    font-family: var(--heading-font); font-weight: 700; font-size: 14px;
    color: var(--accent-2); margin: 14px 0 4px 0;
}}
.shading-label:first-child {{ margin-top: 0; }}

/* ---------- About section ---------- */
.feature-grid {{ display: flex; flex-wrap: wrap; gap: 18px; justify-content: center; margin-top: 12px; }}
/* Primary AI feature cards (Text-to-Image / Photo-to-Sketch / Speech-to-Text /
   Text-to-Text) are always centered as a group and forced to identical
   size/shape so the row reads as one balanced set. */
.feature-grid--primary {{ display: flex; flex-wrap: wrap; gap: 18px; justify-content: center; align-items: stretch; margin-top: 12px; }}
.feature-grid--primary .feature-card {{ flex: 1 1 250px; max-width: 270px; min-height: 250px; display: flex; flex-direction: column; align-items: center; }}
.feature-card {{
    flex: 1 1 260px; max-width: 340px;
    background: {P['card_bg']};
    border: 1px solid {P['panel_border']}; border-radius: 18px; padding: 26px;
    text-align: center;
    transition: background var(--dur-slow) var(--ease), transform var(--dur-med) var(--ease),
                border-color var(--dur-med) var(--ease), box-shadow var(--dur-med) var(--ease);
    animation: fade-in-up .5s var(--ease);
}}
.feature-card:hover {{ transform: translateY(-8px) scale(1.015); border-color: rgba(56,189,248,.45); box-shadow: 0 18px 34px rgba(0,0,0,.25); }}
.feature-icon {{
    display: flex; align-items: center; justify-content: center;
    width: 56px; height: 56px; margin: 0 auto 14px auto; border-radius: 16px;
    background: linear-gradient(135deg, rgba(56,189,248,.16), rgba(37,99,235,.1));
    color: var(--accent-2);
    transition: transform var(--dur-med) var(--ease), color var(--dur-med) var(--ease), background var(--dur-med) var(--ease);
}}
.feature-icon svg {{ width: 28px; height: 28px; stroke: currentColor; }}
.feature-card:hover .feature-icon {{ transform: scale(1.12) rotate(-4deg); }}
.feature-title {{ font-family: var(--heading-font); font-size: 18px; font-weight: 700; color: var(--text-primary); margin-bottom: 8px; letter-spacing: .2px; }}
.feature-text {{ font-size: 14px; color: var(--text-secondary); line-height: 1.5; }}

/* ---------- Developer profile card ---------- */
.profile-wrap {{ display: flex; justify-content: center; margin-top: 12px; }}
.profile-card {{
    display: flex; align-items: center; gap: 26px; flex-wrap: wrap; justify-content: center;
    background: {P['card_bg']};
    border: 1px solid rgba(56,189,248,.25); border-radius: 22px; padding: 30px 40px;
    max-width: 640px;
    transition: background var(--dur-slow) var(--ease), transform var(--dur-med) var(--ease), box-shadow var(--dur-med) var(--ease);
    animation: fade-in-up .5s var(--ease);
}}
.profile-card:hover {{ transform: translateY(-5px); box-shadow: 0 22px 42px rgba(0,0,0,.28); }}
.profile-photo-wrap {{ flex-shrink: 0; }}
.profile-photo {{
    width: 140px; height: 140px; border-radius: 50%; object-fit: cover;
    border: 3px solid var(--accent-2); box-shadow: 0 0 30px rgba(56,189,248,.25);
    transition: transform var(--dur-med) var(--ease), box-shadow var(--dur-med) var(--ease);
}}
.profile-photo:hover {{ transform: scale(1.06) rotate(1deg); box-shadow: 0 0 40px rgba(56,189,248,.4); }}
.profile-info {{ text-align: left; }}
.profile-name {{ font-family: var(--heading-font); font-size: 22px; font-weight: 700; color: var(--text-primary); margin: 0 0 4px 0; }}
.profile-role {{ font-size: 14px; color: var(--accent-2); font-weight: 600; margin: 0 0 10px 0; }}
.profile-detail {{ font-size: 14px; color: var(--text-secondary); margin: 2px 0; display: flex; align-items: center; gap: 8px; }}
.profile-detail-icon {{ display: inline-flex; flex-shrink: 0; }}
.profile-detail-icon svg {{ width: 15px; height: 15px; stroke: var(--accent-2); }}

/* ---------- Native Streamlit chrome (header, toolbar, Deploy button, ⋮ MainMenu) ---------- */
/* These are rendered by Streamlit itself, outside our layout -- themed here so
   they always match the selected palette instead of Streamlit's own default. */
[data-testid="stHeader"] {{
    background: {P['navbar_bg']} !important;
    transition: background var(--dur-slow) var(--ease), box-shadow var(--dur-med) var(--ease);
}}
[data-testid="stToolbar"], [data-testid="stToolbarActions"] {{
    transition: background var(--dur-med) var(--ease);
}}
[data-testid="stToolbar"] button,
[data-testid="stToolbarActions"] button,
[data-testid="stDeployButton"] button,
#MainMenu button {{
    background: {P['navbar_bg']} !important;
    color: var(--text-primary) !important;
    border: 1px solid {P['navbar_border']} !important;
    border-radius: 10px !important;
    transition: background var(--dur-fast) var(--ease), color var(--dur-fast) var(--ease),
                border-color var(--dur-fast) var(--ease), transform var(--dur-fast) var(--ease);
}}
[data-testid="stToolbar"] button:hover,
[data-testid="stToolbarActions"] button:hover,
[data-testid="stDeployButton"] button:hover,
#MainMenu button:hover {{
    background: rgba(56,189,248,.15) !important;
    transform: translateY(-1px);
}}
[data-testid="stToolbar"] svg,
[data-testid="stToolbarActions"] svg,
[data-testid="stDeployButton"] svg,
#MainMenu svg,
#MainMenu span {{
    fill: var(--text-primary) !important;
    color: var(--text-primary) !important;
    transition: fill var(--dur-fast) var(--ease), color var(--dur-fast) var(--ease);
}}

/* ---------- Popovers (Appearance ⋮ menu + mobile ☰ menu) always follow theme ---------- */
div[data-testid="stPopoverBody"] {{
    background: {P['panel_bg']} !important;
    backdrop-filter: blur(14px);
    border: 1px solid {P['panel_border']} !important;
    border-radius: 14px !important;
    box-shadow: 0 18px 40px rgba(0,0,0,.28) !important;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-med) var(--ease);
}}
div[data-testid="stPopoverBody"] * {{ color: var(--text-primary); }}

/* ---------- Prompt input area + all form controls follow theme ---------- */
.stTextArea textarea, .stTextInput input, .stNumberInput input {{
    background: {P['panel_bg']} !important;
    color: var(--text-primary) !important;
    border: 1px solid {P['panel_border']} !important;
    border-radius: 10px !important;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-fast) var(--ease),
                color var(--dur-med) var(--ease), box-shadow var(--dur-fast) var(--ease);
}}
.stTextArea textarea::placeholder, .stTextInput input::placeholder {{ color: var(--text-secondary) !important; opacity: .85; }}
.stTextArea textarea:focus, .stTextInput input:focus, .stNumberInput input:focus {{
    border-color: var(--accent-2) !important;
    box-shadow: 0 0 0 3px rgba(56,189,248,.16) !important;
}}
[data-baseweb="select"] > div, [data-baseweb="base-input"] {{
    background: {P['panel_bg']} !important;
    border-color: {P['panel_border']} !important;
    color: var(--text-primary) !important;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-fast) var(--ease),
                box-shadow var(--dur-fast) var(--ease);
}}
.stSelectbox:hover [data-baseweb="select"] > div {{ border-color: {P['panel_border_hover']} !important; }}
.stSelectbox [data-baseweb="select"]:has(input:focus) > div {{
    border-color: var(--accent-2) !important;
    box-shadow: 0 0 0 3px rgba(56,189,248,.16) !important;
}}
[data-baseweb="popover"] [data-baseweb="menu"], ul[role="listbox"] {{
    background: {P['panel_bg']} !important;
    border: 1px solid {P['panel_border']} !important;
}}
li[role="option"] {{ color: var(--text-primary) !important; transition: background var(--dur-fast) var(--ease); }}
li[role="option"]:hover {{ background: rgba(56,189,248,.15) !important; }}
.stFileUploader section, .stCameraInput video {{
    background: {P['panel_bg']} !important;
    border-color: {P['panel_border']} !important;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-fast) var(--ease);
}}

/* ---------- File uploader dropzone -- icon + size-limit label always on one
   aligned row (upload icon left, "200MB" limit label right), at every
   breakpoint including mobile, instead of Streamlit's default stacked layout
   that could wrap/overlap on narrow screens. ---------- */
[data-testid="stFileUploaderDropzone"] > div {{
    display: flex !important;
    align-items: center !important;
    flex-wrap: nowrap !important;
    gap: 10px;
    width: 100%;
    min-width: 0;
}}
[data-testid="stFileUploaderDropzone"] svg {{ flex-shrink: 0; }}
[data-testid="stFileUploaderDropzoneInstructions"] {{
    display: flex !important;
    flex-direction: row !important;
    align-items: center !important;
    justify-content: space-between !important;
    flex: 1 1 auto;
    min-width: 0;
    gap: 8px;
    width: 100%;
}}
[data-testid="stFileUploaderDropzoneInstructions"] > div {{
    display: flex; flex-direction: row; align-items: center; justify-content: space-between;
    width: 100%; min-width: 0; gap: 8px;
}}
[data-testid="stFileUploaderDropzoneInstructions"] span:last-child,
[data-testid="stFileUploaderDropzoneInstructions"] small {{
    margin-left: auto;
    white-space: nowrap;
    text-align: right;
    flex-shrink: 0;
}}

/* ---------- Expanders, alerts, code blocks, dataframes -- follow theme too ---------- */
[data-testid="stExpander"] {{
    background: {P['panel_bg']} !important;
    border: 1px solid {P['panel_border']} !important;
    border-radius: 14px !important;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-med) var(--ease);
}}
[data-testid="stExpander"] summary {{ color: var(--text-primary) !important; transition: color var(--dur-med) var(--ease); }}
div[data-testid="stAlertContentInfo"], div[data-testid="stAlertContentSuccess"],
div[data-testid="stAlertContentWarning"], div[data-testid="stAlertContentError"],
div[data-testid^="stAlert"] {{
    transition: background var(--dur-med) var(--ease), color var(--dur-med) var(--ease),
                border-color var(--dur-med) var(--ease);
}}
.stCodeBlock, .stCodeBlock pre, div[data-testid="stCode"] {{
    background: {P['panel_bg']} !important;
    border: 1px solid {P['panel_border']} !important;
    border-radius: 12px !important;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-med) var(--ease);
}}
[data-testid="stDataFrame"] {{
    border: 1px solid {P['panel_border']} !important;
    border-radius: 12px !important;
    transition: border-color var(--dur-med) var(--ease);
}}

/* ---------- AI Chatbot module ---------- */
[data-testid="stChatMessage"] {{
    background: {P['panel_bg']} !important;
    border: 1px solid {P['panel_border']} !important;
    border-radius: 16px !important;
    padding: 4px 6px !important;
    margin-bottom: 10px;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-med) var(--ease),
                transform var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease);
    animation: fade-in .35s var(--ease);
}}
[data-testid="stChatMessage"]:hover {{ border-color: {P['panel_border_hover']}; box-shadow: 0 8px 20px rgba(0,0,0,.12); }}
[data-testid="stChatMessageAvatarCustom"], [data-testid="stChatMessageAvatarUser"],
[data-testid="stChatMessageAvatarAssistant"] {{
    transition: transform var(--dur-fast) var(--ease);
}}
[data-testid="stChatMessage"]:hover [data-testid^="stChatMessageAvatar"] {{ transform: scale(1.08); }}
[data-testid="stChatInput"] textarea {{
    background: {P['panel_bg']} !important;
    color: var(--text-primary) !important;
    border-radius: 12px !important;
    transition: background var(--dur-med) var(--ease), border-color var(--dur-fast) var(--ease),
                box-shadow var(--dur-fast) var(--ease);
}}
[data-testid="stChatInput"]:focus-within textarea {{
    border-color: var(--accent-2) !important;
    box-shadow: 0 0 0 3px rgba(56,189,248,.16) !important;
}}
[data-testid="stChatInputSubmitButton"] {{
    transition: transform var(--dur-fast) var(--ease), background var(--dur-fast) var(--ease);
}}
[data-testid="stChatInputSubmitButton"]:hover {{ transform: scale(1.1); }}
[data-testid="stChatInputSubmitButton"]:active {{ transform: scale(.94); }}

/* ---------- Voice-to-Prompt mic trigger ---------- */
.mic-btn-wrap {{ display: flex; align-items: center; gap: 8px; margin: 4px 0 2px 0; }}
.mic-btn {{
    display: inline-flex; align-items: center; justify-content: center;
    width: 40px; height: 40px; border-radius: 50%;
    background: linear-gradient(135deg, var(--accent), #1d4ed8);
    color: #fff; border: none; cursor: pointer; font-size: 18px; padding: 0;
    box-shadow: 0 4px 14px rgba(37,99,235,.35);
    transition: transform var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease), filter var(--dur-fast) var(--ease);
}}
.mic-btn svg {{ width: 19px; height: 19px; stroke: #fff; fill: none; flex-shrink: 0; pointer-events: none; }}
.mic-btn:hover {{ transform: scale(1.08); box-shadow: 0 8px 20px rgba(37,99,235,.45); filter: brightness(1.08); }}
.mic-btn:active {{ transform: scale(.95); }}
.mic-btn.listening {{ animation: mic-pulse 1.1s ease-in-out infinite; background: linear-gradient(135deg, #ef4444, #dc2626); }}
@keyframes mic-pulse {{
    0%   {{ box-shadow: 0 0 0 0 rgba(239,68,68,.55); }}
    70%  {{ box-shadow: 0 0 0 12px rgba(239,68,68,0); }}
    100% {{ box-shadow: 0 0 0 0 rgba(239,68,68,0); }}
}}
.mic-status {{ font-size: 12.5px; color: var(--text-secondary); transition: color var(--dur-med) var(--ease); }}

/* ---------- Studio block (the 4 AI modules) -- original side-by-side
   dashboard layout: control panel on the left, output/preview panel on the
   right, centered as a balanced unit on desktop. ---------- */
.st-key-studio_columns {{
    max-width: 1180px;
    margin-left: auto;
    margin-right: auto;
}}

/* ---------- Small screens ---------- */
@media (max-width: 768px) {{
    div[data-testid="stRadio"] label {{ padding: 8px 16px; font-size: 14px; }}
    [data-testid="stVerticalBlockBorderWrapper"] {{ padding: 14px !important; }}
    .profile-card {{ flex-direction: column; text-align: center; padding: 24px; }}
    .profile-info {{ text-align: center; width: 100%; }}
    .profile-detail {{ justify-content: center; }}
}}
@media (max-width: 640px) {{
    .density-panel {{ height: auto; padding-right: 0; margin-bottom: 10px; }}
    .shading-panel {{ height: auto; padding-right: 0; border-right: none; border-bottom: 1px dashed {P['panel_border']}; padding-bottom: 12px; margin-bottom: 12px; text-align: center; }}
    /* The hamburger menu (now also holding Appearance) stays pinned top-right
       and comfortably tappable on small devices. */
    .st-key-nav_mobile {{ margin-right: 0; display: flex !important; }}
    .st-key-nav_mobile button {{ padding: 6px 11px !important; font-size: 16px !important; }}

    /* ---- Center-align primary content sections on mobile ----
       Streamlit already stacks st.columns() vertically below this
       breakpoint; here we just center the content within each stacked
       card so text, headers, and controls read as balanced, not stacking
       logic (which Streamlit itself handles). */
    .st-key-studio_columns [data-testid="column"] {{
        display: flex; flex-direction: column; align-items: center;
        margin-bottom: 22px;
    }}
    .st-key-studio_columns [data-testid="column"]:last-child {{ margin-bottom: 0; }}
    .st-key-studio_columns {{ text-align: center; }}
    .st-key-studio_columns .panel-header {{ justify-content: center; text-align: center; border-left: none; border-top: 3px solid var(--accent); padding: 8px 0 0 0; width: 100%; }}
    .st-key-studio_columns .stButton, .st-key-studio_columns .stDownloadButton,
    .st-key-studio_columns .stTextArea, .st-key-studio_columns .stSelectbox,
    .st-key-studio_columns .stRadio, .st-key-studio_columns .stFileUploader,
    .st-key-studio_columns .stCameraInput, .st-key-studio_columns .stExpander,
    .st-key-studio_columns .stCode,
    .st-key-studio_columns [data-testid="stVerticalBlockBorderWrapper"] {{ width: 100%; }}
    .st-key-studio_columns .stButton, .st-key-studio_columns .stDownloadButton {{ display: flex; justify-content: center; }}
    .st-key-studio_columns .mic-btn-wrap {{ justify-content: center; }}
    .mode-switch-wrap div[data-testid="stRadio"] > div[role="radiogroup"] {{
        width: 100%; flex-direction: column; flex-wrap: nowrap; gap: 10px;
        background: transparent; border: none; padding: 0;
    }}
    .mode-switch-wrap div[data-testid="stRadio"] label {{
        display: flex; align-items: center; justify-content: center; gap: 10px;
        width: 100%; padding: 15px 18px; font-size: 15.5px;
        background: {P['radio_bg']}; border: 1px solid rgba(56,189,248,.25); border-radius: 14px;
    }}

    /* File uploader on mobile: icon pinned left, "200MB" size-limit label
       pinned right, both on the same line -- the longer "Drag and drop file
       here" caption is hidden since there's no room for three items on one
       row on narrow screens and the limit label is the more useful bit. */
    [data-testid="stFileUploaderDropzone"] {{ padding: 12px 14px !important; }}
    [data-testid="stFileUploaderDropzoneInstructions"] span:first-child {{ display: none; }}
    [data-testid="stFileUploaderDropzoneInstructions"] span:last-child,
    [data-testid="stFileUploaderDropzoneInstructions"] small {{
        font-size: 11.5px !important;
        overflow: hidden;
        text-overflow: ellipsis;
        max-width: 100%;
    }}
}}
@media (max-width: 380px) {{
    .st-key-nav_mobile button, .st-key-theme_menu_wrap button {{ padding: 5px 9px !important; font-size: 14px !important; }}
    .nav-title {{ max-width: 46vw; }}
}}

/* ---------- Extra micro-interactions & polish ---------- */
[data-baseweb="select"] > div {{ transition: background var(--dur-med) var(--ease), border-color var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease); }}
[data-baseweb="select"] > div:hover {{ border-color: {P['panel_border_hover']} !important; }}
.stTextArea textarea:hover, .stTextInput input:hover, .stNumberInput input:hover {{ border-color: {P['panel_border_hover']} !important; }}
div[data-testid^="stAlert"] {{ animation: fade-in-up .35s var(--ease); }}
[data-testid="stExpander"] summary {{ transition: color var(--dur-fast) var(--ease), background var(--dur-fast) var(--ease); }}
[data-testid="stExpander"] summary:hover {{ color: var(--accent-2) !important; }}
.stSlider [role="slider"] {{ transition: transform var(--dur-fast) var(--ease), box-shadow var(--dur-fast) var(--ease); }}
.stSlider [role="slider"]:hover {{ transform: scale(1.15); }}
[data-testid="stFileUploaderDropzone"], .stFileUploader section {{ transition: border-color var(--dur-fast) var(--ease), background var(--dur-fast) var(--ease), transform var(--dur-fast) var(--ease); }}
[data-testid="stFileUploaderDropzone"]:hover, .stFileUploader section:hover {{ border-color: var(--accent-2) !important; transform: translateY(-1px); }}
[data-testid="stDataFrame"] {{ transition: box-shadow var(--dur-med) var(--ease); }}
[data-testid="stDataFrame"]:hover {{ box-shadow: 0 10px 24px rgba(0,0,0,.14); }}
</style>
""", unsafe_allow_html=True)

# ------------------------------------------------------------------
# CONSTANTS
# ------------------------------------------------------------------
PROMPT_PRESETS = {
    "Custom Prompt": "A majestic royal Bengal tiger sitting in deep white snow, looking at the camera",
    "Premium VIP Event Ticket": "A premium VIP event ticket design, sleek black background with golden metallic neon lines, modern typography",
    "Cyberpunk Neon City": "A futuristic cyberpunk city street at night, glowing neon signs in blue and magenta, light rain reflections",
    "Studio Portrait of a Lion": "A majestic male African lion with a massive dark mane, sitting proudly on a rock at sunset, studio lighting",
}
STYLE_PRESETS = [
    "Photorealistic Cinematic Lighting",
    "National Geographic Documentary Style",
    "Studio Portrait Photography",
    "3D Digital Anime Art",
]
TEXT_PRESETS = {
    "Custom Prompt": "Write a 400-word assignment on the causes and effects of climate change.",
    "Programming Task": "Write a function that checks whether a given string is a palindrome, with comments explaining each step.",
}
CODE_LANGUAGES = ["C++", "Python", "Java", "JavaScript", "SQL", "C", "C#", "TypeScript"]
CODE_FILE_EXT = {
    "Python": "py", "C++": "cpp", "Java": "java", "JavaScript": "js",
    "C": "c", "C#": "cs", "TypeScript": "ts", "SQL": "sql",
}
CODE_HIGHLIGHT_LANG = {
    "Python": "python", "C++": "cpp", "Java": "java", "JavaScript": "javascript",
    "C": "c", "C#": "csharp", "TypeScript": "typescript", "SQL": "sql",
}
SPEECH_LANGUAGES = {"English": "en-US", "Urdu": "ur-PK"}
VIDEO_PROMPT_PRESETS = {
    "Custom Prompt": "A majestic royal Bengal tiger walking slowly through deep white snow, cinematic camera pan",
    "Cyberpunk Street": "A futuristic cyberpunk city street at night, glowing neon signs, light rain, camera slowly pushing forward",
    "Ocean Sunset": "A calm ocean at sunset with gentle waves rolling in, golden light reflecting on the water",
    "Paper Airplane": "A paper airplane gliding gracefully through a sunlit office, slow motion, dust particles in the light",
}
# Curated subset of Pollinations.ai's free-tier-friendly video models,
# mapped to their underlying model id. "Fast & Reliable" (zimage) is listed
# first -- and used as the default selection -- since it is confirmed to
# work end-to-end on the free anonymous tier; the others remain available
# for anyone with a configured API key that unlocks them.
VIDEO_MODELS = {
    "Fast & Reliable": "zimage",
    "Balanced (Fast)": "wan-fast",
    "Cinematic": "veo",
    "High Detail": "seedance-pro",
}
VIDEO_ASPECTS = {"Widescreen (16:9)": "16:9", "Vertical (9:16)": "9:16"}
VIDEO_DURATION_RANGE = (2, 10)
ASPECT_RATIOS = {
    "Square (512x512)": (512, 512),
    "Landscape (768x512)": (768, 512),
    "Portrait (512x768)": (512, 768),
    "Ultra HD Square (1024x1024)": (1024, 1024),
}
# (label, page_key, material icon) -- clean, modern glyphs shown via Streamlit's
# native `icon=` button parameter instead of emoji.
NAV_ITEMS = [
    ("Home", "home", ":material/home:"),
    ("About", "about", ":material/info:"),
    ("Developer", "dev", ":material/code:"),
]
# Professional Material Icon labels for the four AI modules -- shared
# constants so the mode switch, and every place that programmatically
# selects a mode, always agree on the exact same label text.
MODE_TEXT_TO_IMAGE = ":material/image: Text-to-Image"
MODE_PHOTO_TO_SKETCH = ":material/draw: Photo-to-Sketch"
MODE_SPEECH_TO_TEXT = ":material/mic: Speech-to-Text"
MODE_TEXT_TO_TEXT = ":material/text_fields: Text-to-Text"
MODE_TEXT_TO_VIDEO = ":material/movie: Text-to-Video"
MODE_AI_CHATBOT = ":material/forum: AI Chatbot"
DOWNLOAD_SIZES = {"Small": 480, "Medium": 960, "Large": 1600}
DEVELOPER = {
    "name": "Absar Ali",
    "degree": "BS Computer Science",
    "profession": "Software Developer",
    # Drop a photo at any of these paths (relative to this script) to display it.
    "photo_candidates": [
        "assets/developer.jpg", "assets/developer.jpeg", "assets/developer.png",
        "developer.jpg", "developer.jpeg", "developer.png",
    ],
}

# Shared, professional icon set (consistent stroke width/size/style) used for
# both panel headers and the About page feature cards -- a single source so
# every icon across the app shares the same visual language.
ICONS = {
    "image": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <rect x="3" y="4" width="18" height="16" rx="2.5"></rect>
        <circle cx="8.5" cy="9.5" r="1.6"></circle>
        <path d="M21 16l-5.5-5.5a1.5 1.5 0 0 0-2.1 0L4 19"></path>
    </svg>""",
    "pencil": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M4 20l4-1 11-11a2 2 0 0 0 0-2.8l-.2-.2a2 2 0 0 0-2.8 0L5 16l-1 4z"></path>
        <path d="M14 6.5l3 3"></path>
    </svg>""",
    "mic": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <rect x="9" y="3" width="6" height="11" rx="3"></rect>
        <path d="M5 11a7 7 0 0 0 14 0"></path>
        <path d="M12 18v3"></path>
        <path d="M9 21h6"></path>
    </svg>""",
    "text": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M4 5h16"></path>
        <path d="M4 10h16"></path>
        <path d="M4 15h10"></path>
        <path d="M4 20h6"></path>
    </svg>""",
    "bolt": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M13 2 4 14h6l-1 8 9-12h-6l1-8z"></path>
    </svg>""",
    "settings": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <circle cx="12" cy="12" r="3.2"></circle>
        <path d="M19.4 13.5a1.7 1.7 0 0 0 .34 1.87l.06.06a2 2 0 1 1-2.83 2.83l-.06-.06a1.7 1.7 0 0 0-1.87-.34 1.7 1.7 0 0 0-1.04 1.56V19.9a2 2 0 1 1-4 0v-.09a1.7 1.7 0 0 0-1.11-1.56 1.7 1.7 0 0 0-1.87.34l-.06.06a2 2 0 1 1-2.83-2.83l.06-.06a1.7 1.7 0 0 0 .34-1.87 1.7 1.7 0 0 0-1.56-1.04H4.1a2 2 0 1 1 0-4h.09a1.7 1.7 0 0 0 1.56-1.11 1.7 1.7 0 0 0-.34-1.87l-.06-.06a2 2 0 1 1 2.83-2.83l.06.06a1.7 1.7 0 0 0 1.87.34H10.2a1.7 1.7 0 0 0 1.04-1.56V4.1a2 2 0 1 1 4 0v.09a1.7 1.7 0 0 0 1.04 1.56 1.7 1.7 0 0 0 1.87-.34l.06-.06a2 2 0 1 1 2.83 2.83l-.06.06a1.7 1.7 0 0 0-.34 1.87V10.2a1.7 1.7 0 0 0 1.56 1.04h.09a2 2 0 1 1 0 4h-.09a1.7 1.7 0 0 0-1.56 1.04z"></path>
    </svg>""",
    "download": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M12 3v12"></path>
        <path d="M6.5 10.5 12 16l5.5-5.5"></path>
        <path d="M4 20h16"></path>
    </svg>""",
    "camera": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M4 8.5A1.5 1.5 0 0 1 5.5 7H8l1.2-2h5.6L16 7h2.5A1.5 1.5 0 0 1 20 8.5v9A1.5 1.5 0 0 1 18.5 19h-13A1.5 1.5 0 0 1 4 17.5z"></path>
        <circle cx="12" cy="12.5" r="3.4"></circle>
    </svg>""",
    "info": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <circle cx="12" cy="12" r="9"></circle>
        <path d="M12 11v6"></path>
        <circle cx="12" cy="7.6" r="0.9" fill="currentColor" stroke="none"></circle>
    </svg>""",
    "developer": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <rect x="3" y="4.5" width="18" height="13" rx="2"></rect>
        <path d="M8 9.5 5.5 12 8 14.5"></path>
        <path d="M16 9.5 18.5 12 16 14.5"></path>
        <path d="M13 8.5 11 15.5"></path>
        <path d="M8.5 20.5h7"></path>
        <path d="M12 17.5v3"></path>
    </svg>""",
    "graduation": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M2.5 8.5 12 4l9.5 4.5L12 13z"></path>
        <path d="M6.5 10.6v4.4c0 1.4 2.5 3 5.5 3s5.5-1.6 5.5-3v-4.4"></path>
        <path d="M21.5 8.5v6"></path>
    </svg>""",
    "idea": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M9 18h6"></path>
        <path d="M10 21h4"></path>
        <path d="M12 3a6.5 6.5 0 0 0-3.7 11.8c.6.44.95 1.13.95 1.87V17h5.5v-.33c0-.74.35-1.43.95-1.87A6.5 6.5 0 0 0 12 3z"></path>
    </svg>""",
    "chat": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <path d="M4 5.5A2.5 2.5 0 0 1 6.5 3h11A2.5 2.5 0 0 1 20 5.5v8A2.5 2.5 0 0 1 17.5 16H10l-4.5 4v-4H6.5A2.5 2.5 0 0 1 4 13.5z"></path>
        <path d="M8 8.2h8"></path>
        <path d="M8 11.4h5"></path>
    </svg>""",
    "video": """<svg viewBox="0 0 24 24" fill="none" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
        <rect x="3" y="6" width="13" height="12" rx="2.2"></rect>
        <path d="M16 10.3 21 7.5v9l-5-2.8z"></path>
    </svg>""",
}


def brand_logo_svg(uid: str, size: int = 28) -> str:
    """Modern, premium, AI-inspired brand mark -- a glowing neural-node
    cluster rendered as a self-contained gradient SVG (independent of the
    surrounding text color so it always reads clearly in both themes).
    `uid` keeps the gradient id unique when the mark is rendered more than
    once on the same page (navbar + hero)."""
    gid = f"ncLogoGrad-{uid}"
    return f"""<svg width="{size}" height="{size}" viewBox="0 0 32 32" role="img" aria-label="NeuralCraft logo">
        <defs>
            <linearGradient id="{gid}" x1="0%" y1="0%" x2="100%" y2="100%">
                <stop offset="0%" stop-color="#7dd3fc"/>
                <stop offset="55%" stop-color="#38bdf8"/>
                <stop offset="100%" stop-color="#7c3aed"/>
            </linearGradient>
        </defs>
        <circle cx="16" cy="16" r="15" fill="url(#{gid})" opacity="0.14"/>
        <g fill="none" stroke="url(#{gid})" stroke-width="1.7" stroke-linecap="round" stroke-linejoin="round">
            <path d="M16 9.6 16 13.6 M14.1 17.6 9.4 20.4 M17.9 17.6 22.6 20.4"/>
        </g>
        <g fill="url(#{gid})">
            <circle cx="16" cy="7.2" r="2.3"/>
            <circle cx="7.6" cy="22.4" r="2.3"/>
            <circle cx="24.4" cy="22.4" r="2.3"/>
            <circle cx="16" cy="16" r="2.5"/>
        </g>
    </svg>"""


def icon_header(icon_key: str, label: str, style: str = ""):
    """Renders a `.panel-header` with a small inline-SVG icon in front of the
    label -- the shared, single source of truth for panel-header markup so
    every header across the app uses the exact same icon set, size, and
    spacing instead of ad-hoc emoji."""
    st.markdown(
        f'<div class="panel-header" style="{style}">'
        f'<span class="panel-header-icon">{ICONS[icon_key]}</span>{label}</div>',
        unsafe_allow_html=True,
    )


# ------------------------------------------------------------------
# CORE ENGINES
# ------------------------------------------------------------------
def polish_generated_image(image_bytes: bytes, target_width: int = None, target_height: int = None) -> bytes:
    """Post-processing polish applied to freshly generated images: the image
    is rendered at a higher-than-requested resolution upstream and then
    Lanczos-downsampled back to the target size (supersampling), which
    smooths jagged edges and boosts perceived detail; a final detail-enhance
    and unsharp pass then makes the result read as sharper, higher-resolution,
    and more natural/realistic without altering the composition."""
    try:
        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    except Exception:
        return image_bytes
    if target_width and target_height and (img.width != target_width or img.height != target_height):
        img = img.resize((target_width, target_height), Image.LANCZOS)
    img = img.filter(ImageFilter.DETAIL)
    img = ImageEnhance.Sharpness(img).enhance(1.32)
    img = ImageEnhance.Contrast(img).enhance(1.07)
    img = ImageEnhance.Color(img).enhance(1.05)
    img = img.filter(ImageFilter.UnsharpMask(radius=1.6, percent=120, threshold=2))
    buf = io.BytesIO()
    img.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


# Single, user-friendly message shown for any failure caused by the user's
# own connectivity (no internet, DNS failure, connection dropped, request
# timed out) -- used consistently across every AI feature instead of raw
# exception text or generic server-error wording.
NETWORK_ERROR_MESSAGE = "Network Error. Please check your internet connection and try again."


def generate_ai_image(prompt: str, width: int, height: int, style: str, progress_cb=None):
    """Free text-to-image generation via the Pollinations.ai 'flux' model.
    Renders at a supersampled resolution (1.5x, capped) for better prompt
    understanding and finer upstream detail, then downsamples back to the
    requested size during polishing for sharper, higher-resolution, more
    realistic and natural-looking results.

    Retries transient connectivity hiccups and momentary upstream overload
    with backoff (mirroring generate_ai_text) instead of failing on the
    first bad response, so generation feels reliable under flaky networks.
    `progress_cb(attempt, max_attempts, message)` optionally surfaces retry
    status to the UI.
    """
    full_prompt = (
        f"{prompt}, {style}, ultra-detailed, high resolution, 8k, sharp focus, "
        f"intricate details, professional quality, masterpiece, hyper-realistic lighting, "
        f"natural skin and material textures, accurate anatomy and proportions, "
        f"balanced exposure, true-to-life color grading"
    )
    encoded = urllib.parse.quote(full_prompt)
    seed = random.randint(1000, 99999)
    scale = 1.5
    render_w = min(int(width * scale), 1536)
    render_h = min(int(height * scale), 1536)
    url = (
        f"https://image.pollinations.ai/prompt/{encoded}"
        f"?width={render_w}&height={render_h}&seed={seed}&model=flux&nologo=true&enhance=true"
    )

    max_attempts = 3
    last_error = NETWORK_ERROR_MESSAGE
    for attempt in range(max_attempts):
        if progress_cb and attempt > 0:
            progress_cb(attempt + 1, max_attempts, last_error)
        try:
            response = requests.get(url, timeout=45)
        except requests.RequestException:
            # Covers connection errors, timeouts, DNS failures, dropped
            # connections -- all surfaced with the same clear message.
            last_error = NETWORK_ERROR_MESSAGE
            if attempt < max_attempts - 1:
                time.sleep(min(1.5 * (attempt + 1), 5))
            continue

        if response.status_code == 200:
            return polish_generated_image(response.content, width, height), None
        elif response.status_code in _RETRYABLE_STATUS or response.status_code == 429:
            last_error = f"The AI service is temporarily unavailable (status {response.status_code}). Retrying..."
            if attempt < max_attempts - 1:
                time.sleep(min(2 * (attempt + 1), 6))
            continue
        else:
            return None, f"Server returned status {response.status_code}. Please retry."

    return None, last_error


def _strip_code_fences(text: str) -> str:
    """Extracts just the code from inside the first fenced ``` code block
    in the generated text (with or without a language tag), discarding any
    explanatory text the model wrote before or after the fence -- this is
    what previously broke the code display/download when the model added a
    short note after the closing fence. Falls back to the plain stripped
    text if no fenced block is present, so unfenced responses still work."""
    match = re.search(r"```[a-zA-Z0-9_+-]*\r?\n(.*?)```", text, re.DOTALL)
    if match:
        return match.group(1).strip("\n").rstrip()
    return text.strip()


# Pollinations has, over time, hosted its OpenAI-compatible chat endpoint on
# more than one subdomain; gen.pollinations.ai is the current, actively
# maintained host, with the older text.pollinations.ai host kept as a
# compatible fallback for the rare case the primary host itself is down.
TEXT_ENDPOINT_PRIMARY = "https://gen.pollinations.ai/v1/chat/completions"
TEXT_ENDPOINT_FALLBACK = "https://text.pollinations.ai/openai"

# Server-side errors that are transient (upstream overloaded / restarting /
# briefly unreachable) -- worth trying the fallback host for, since retrying
# the *same* host repeatedly is what made the module feel unreliable before.
_RETRYABLE_STATUS = {429, 500, 502, 503, 504}


def _build_system_prompt(content_type: str, language: str = None) -> str:
    if content_type == "Programming Code" and language:
        return (
            f"You are an expert {language} software engineer. Write clean, correct, "
            f"efficient, well-commented {language} code that solves the task exactly. "
            f"Return ONLY a single fenced code block containing the complete, runnable "
            f"code -- no text before or after the code block."
        )
    return (
        "You are an expert academic writer and patient teacher. Produce accurate, "
        "clearly organized, high-quality written content tailored to the request -- "
        "whether it's an assignment, report, explanation, or general content -- "
        "with a logical structure (introduction, well-organized body, conclusion) "
        "and headings where useful. Do not artificially shorten or summarize the "
        "response -- write it out in full, comprehensive detail, using as much "
        "length as the topic genuinely calls for."
    )


def _extract_chat_content(response: requests.Response) -> str:
    """Pulls the generated text out of either the standard OpenAI-style JSON
    envelope or a plain-text legacy response, whichever the endpoint returns."""
    try:
        data = response.json()
        return (data["choices"][0]["message"]["content"] or "").strip()
    except (ValueError, KeyError, IndexError, TypeError):
        return response.text.strip()


def generate_ai_text(prompt: str, content_type: str, language: str = None):
    """Free text-to-text generation via the Pollinations.ai text API,
    contextually tuned per content type -- Theory (assignments, reports,
    explanations, and general written content) or Programming Code -- for
    accurate, well-structured, high-quality output.

    Uses a plain GET request against text.pollinations.ai with the full
    system+task prompt URL-encoded directly into the path (mirroring the
    same request shape already used for image generation) instead of the
    JSON chat-completions endpoint, which was returning errors. A
    connectivity failure (no internet, DNS failure, dropped connection,
    timeout) is surfaced immediately as a single, clear message rather than
    retried automatically -- the user decides whether and when to try again.
    """
    system_ctx = _build_system_prompt(content_type, language)
    full_prompt = f"{system_ctx}\n\nTask: {prompt.strip()}"
    encoded = urllib.parse.quote(full_prompt)
    seed = random.randint(1000, 99999)
    url = f"https://text.pollinations.ai/{encoded}?model=openai&seed={seed}&private=true"

    try:
        # Long timeout so very long-form generations have room to
        # complete instead of being cut off by a tight network deadline.
        response = requests.get(url, timeout=120)
    except requests.RequestException:
        # Covers connection errors, DNS failures, dropped connections,
        # and timeouts -- shown once, with no automatic retry.
        return None, NETWORK_ERROR_MESSAGE

    if response.status_code == 200 and response.text.strip():
        return response.text.strip(), None
    if response.status_code in _RETRYABLE_STATUS:
        return None, "The AI service is temporarily unavailable. Please try again in a moment."
    return None, f"Server returned status {response.status_code}. Please try again."


CHATBOT_SYSTEM_PROMPT = (
    "You are the NeuralCraft AI Assistant, a helpful, friendly, and knowledgeable "
    "chatbot embedded in the Text-to-Text module. Answer questions clearly and "
    "accurately, keep replies conversational and reasonably concise, and use "
    "markdown formatting (like code blocks or lists) when it improves clarity."
)


def generate_chat_response(chat_messages: list):
    """Conversational AI response for the Text-to-Text module's built-in AI
    Chatbot. Uses the same reliable GET request against text.pollinations.ai
    as generate_ai_text (the JSON chat-completions endpoint this used to
    call was returning errors), with the running conversation folded into a
    single transcript so multi-turn context is preserved."""
    transcript = "\n".join(
        f"{'User' if msg['role'] == 'user' else 'Assistant'}: {msg['content']}"
        for msg in chat_messages
    )
    full_prompt = (
        f"{CHATBOT_SYSTEM_PROMPT}\n\n"
        f"Conversation so far:\n{transcript}\n\n"
        f"Respond as the Assistant to the latest User message above."
    )
    encoded = urllib.parse.quote(full_prompt)
    seed = random.randint(1000, 99999)
    url = f"https://text.pollinations.ai/{encoded}?model=openai&seed={seed}&private=true"

    try:
        response = requests.get(url, timeout=60)
    except requests.RequestException:
        # Covers connection errors, DNS failures, dropped connections,
        # and timeouts -- shown once, with no automatic retry.
        return None, NETWORK_ERROR_MESSAGE

    if response.status_code == 200 and response.text.strip():
        return response.text.strip(), None
    if response.status_code in _RETRYABLE_STATUS:
        return None, "The AI service is temporarily unavailable. Please try again in a moment."
    return None, f"Server returned status {response.status_code}. Please try again."


# Endpoint for the unified Pollinations.ai video generation service.
VIDEO_ENDPOINT = "https://gen.pollinations.ai/video"

# Video generation has no server-reported progress, so this constant maps
# elapsed wait time to a capped 0-92% visual estimate for the progress bar --
# it only ever reaches 100% once the response has actually arrived.
_VIDEO_PROGRESS_ESTIMATE_SECS = 70

# A "successful" response smaller than this is almost certainly an error
# page or empty stub rather than real video -- treated as a failure so the
# user never sees a broken/unplayable download.
_MIN_VALID_VIDEO_BYTES = 2000


# Fallback bearer token used when no `POLLINATIONS_API_KEY` is configured.
# The `gen.pollinations.ai` video endpoint requires *some* Authorization
# header to be present to serve the free anonymous tier -- a request sent
# with no header at all is rejected outright, even though the tier itself
# needs no real credential. This keeps the module working out of the box.
_ANONYMOUS_VIDEO_TOKEN = "5"


def _get_pollinations_api_key():
    """Resolves the Pollinations.ai bearer token used for video generation.

    Checked first in Streamlit secrets (`POLLINATIONS_API_KEY`), then the
    environment, so a real key -- which unlocks the higher-quality metered
    models -- is used automatically whenever one is configured. If none is
    set, falls back to the free anonymous-tier token so the Authorization
    header is always sent and the "Fast & Reliable" model keeps working
    with zero configuration.
    """
    try:
        key = st.secrets.get("POLLINATIONS_API_KEY")
        if key:
            return str(key).strip() or _ANONYMOUS_VIDEO_TOKEN
    except Exception:
        pass
    return os.environ.get("POLLINATIONS_API_KEY", "").strip() or _ANONYMOUS_VIDEO_TOKEN


def generate_ai_video(prompt: str, duration: int, aspect_ratio: str, model_label: str, progress_cb=None):
    """Text-to-video generation via the Pollinations.ai unified 'gen' video
    API (GET /video/{prompt}), sending a Bearer Authorization header on
    every request -- required by the endpoint even for the free anonymous
    tier (see `_get_pollinations_api_key`). Mirrors the retry/backoff shape
    already used by generate_ai_image so transient upstream hiccups are
    retried automatically, and maps every failure mode (auth rejection,
    insufficient service balance, unknown model, rate limiting, transient
    server errors, connectivity failure, or a truncated/invalid body) to
    its own clear, actionable message instead of a raw error.

    The request itself is dispatched on a background thread so the caller
    can keep polling `progress_cb(fraction, message)` with a smooth,
    honest-feeling progress estimate while waiting -- `fraction` is a 0-1
    estimate since the API reports no real progress of its own.
    """
    full_prompt = (
        f"{prompt.strip()}, cinematic motion, smooth camera movement, coherent motion, "
        f"high quality, detailed, natural lighting"
    )
    encoded = urllib.parse.quote(full_prompt)
    seed = random.randint(1000, 99999)
    model_id = VIDEO_MODELS.get(model_label, "zimage")
    # Always send an Authorization header -- the endpoint rejects requests
    # with none at all, even on the free anonymous tier.
    headers = {"Authorization": f"Bearer {_get_pollinations_api_key()}"}
    url = (
        f"{VIDEO_ENDPOINT}/{encoded}"
        f"?model={model_id}&duration={duration}&aspectRatio={aspect_ratio}&seed={seed}"
    )

    max_attempts = 3
    last_error = NETWORK_ERROR_MESSAGE
    for attempt in range(max_attempts):
        if progress_cb and attempt > 0:
            progress_cb(0.05, f"Attempt {attempt + 1}/{max_attempts} -- {last_error}")

        result = {}

        def _worker():
            try:
                result["response"] = requests.get(url, headers=headers, timeout=280)
            except requests.RequestException:
                # Covers connection errors, DNS failures, dropped
                # connections, and timeouts.
                result["error"] = NETWORK_ERROR_MESSAGE

        thread = threading.Thread(target=_worker, daemon=True)
        start = time.time()
        thread.start()
        while thread.is_alive():
            if progress_cb:
                elapsed = time.time() - start
                fraction = min(0.92, elapsed / _VIDEO_PROGRESS_ESTIMATE_SECS)
                progress_cb(fraction, "Rendering your video -- this can take up to a minute or two...")
            time.sleep(0.3)
        thread.join()

        if "error" in result:
            last_error = result["error"]
            if attempt < max_attempts - 1:
                time.sleep(min(2 * (attempt + 1), 6))
            continue

        response = result.get("response")
        if response is None:
            last_error = NETWORK_ERROR_MESSAGE
            if attempt < max_attempts - 1:
                time.sleep(2)
            continue

        content_type = response.headers.get("content-type", "")
        if (response.status_code == 200 and len(response.content) >= _MIN_VALID_VIDEO_BYTES
                and "video" in content_type):
            if progress_cb:
                progress_cb(1.0, "Finalizing...")
            return response.content, None
        elif response.status_code in (401, 403):
            return None, "Video generation couldn't be authorized right now. Please try again in a moment."
        elif response.status_code == 402:
            return None, ("The video service account has insufficient balance to "
                           "generate this video right now. Please try again later.")
        elif response.status_code == 404:
            return None, "The selected video model is currently unavailable. Try a different model."
        elif response.status_code in _RETRYABLE_STATUS:
            last_error = f"The AI video service is temporarily unavailable (status {response.status_code}). Retrying..."
            if attempt < max_attempts - 1:
                time.sleep(min(2 * (attempt + 1), 8))
            continue
        elif response.status_code == 200:
            # A 200 without a usable video body -- treat as a transient
            # upstream hiccup rather than silently returning empty/broken bytes.
            last_error = "The AI service returned an unexpected or incomplete response. Retrying..."
            if attempt < max_attempts - 1:
                time.sleep(2)
            continue
        else:
            return None, f"Server returned status {response.status_code}. Please retry."

    return None, last_error


def to_pencil_sketch(pil_image: Image.Image, blur: int, shade: float, contrast: float) -> Image.Image:
    """Converts a PIL image into a crisp, professional pencil-sketch rendering.

    Pipeline: edge-preserving smoothing -> classic dodge-blend base shading ->
    a resolution-adaptive, denoised adaptive-threshold line layer for sharp,
    hand-inked edges -> the two are combined for natural shading with crisp,
    clean detail, then lightly sharpened for a polished, portfolio-ready finish.
    """
    bgr = cv2.cvtColor(np.array(pil_image.convert("RGB")), cv2.COLOR_RGB2BGR)
    h, w = bgr.shape[:2]

    # Edge-preserving smoothing removes sensor noise/texture while keeping
    # true edges intact, so shading stays clean without muddying detail.
    smoothed = cv2.bilateralFilter(bgr, d=9, sigmaColor=75, sigmaSpace=75)
    gray = cv2.cvtColor(smoothed, cv2.COLOR_BGR2GRAY)

    # Local-contrast boost (CLAHE) before shading so detail in shadows and
    # highlights survives the dodge blend instead of flattening out.
    clahe = cv2.createCLAHE(clipLimit=2.2, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # Classic "dodge blend" pencil base -- soft, natural graphite shading.
    inverted = 255 - gray
    kernel = blur if blur % 2 else blur + 1
    blurred = cv2.GaussianBlur(inverted, (kernel, kernel), 0)
    dodge = cv2.divide(gray, 255 - blurred, scale=256.0)

    # Dual line layer: adaptive threshold for deliberate hand-inked strokes,
    # blended with Canny for finer, cleaner detail edges -- together they hold
    # onto more structure than either method alone. The adaptive-threshold
    # block size scales with image resolution so lines stay clean and
    # proportionate instead of turning noisy/busy on higher-res photos.
    block_size = max(5, (min(h, w) // 180) | 1)
    median = cv2.medianBlur(gray, 5)
    edges_adaptive = cv2.adaptiveThreshold(
        median, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, block_size, 5,
    )
    edges_canny = cv2.bitwise_not(cv2.Canny(median, 40, 130))
    edges = cv2.bitwise_and(edges_adaptive, edges_canny)
    # Open then close the line layer: opening clears isolated noise specks
    # for cleaner edges, closing seals tiny gaps so strokes read as
    # continuous, deliberate pencil lines rather than broken fragments.
    denoise_kernel = np.ones((2, 2), np.uint8)
    edges = cv2.morphologyEx(edges, cv2.MORPH_OPEN, denoise_kernel)
    edges = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, denoise_kernel)

    # Multiply-blend the soft shading with the crisp line layer: sharper
    # edges and better preserved detail than dodge-blend alone.
    combined = (dodge.astype(np.float32) / 255.0) * (edges.astype(np.float32) / 255.0)
    combined = np.clip(combined * 255.0, 0, 255).astype(np.uint8)

    result = Image.fromarray(combined)
    if shade != 1.0:
        result = ImageEnhance.Brightness(result).enhance(1.0 - shade * 0.2)
    if contrast != 1.0:
        result = ImageEnhance.Contrast(result).enhance(contrast)

    # Final unsharp mask for crisp, professional-grade stroke definition.
    result = result.filter(ImageFilter.UnsharpMask(radius=2.2, percent=150, threshold=2))
    return result


def resize_for_download(pil_image: Image.Image, target_width: int) -> Image.Image:
    """Exports an image at the exact selected resolution (Small/Medium/Large),
    both downscaling and upscaling as needed, while preserving quality:
    high-quality LANCZOS resampling, plus a light unsharp pass after any
    upscale to counter softness and keep edges crisp."""
    w, h = pil_image.size
    if w == target_width:
        return pil_image
    ratio = target_width / float(w)
    target_h = max(1, int(round(h * ratio)))
    resized = pil_image.resize((target_width, target_h), Image.LANCZOS)
    if target_width > w:
        # Upscaling can soften fine linework -- restore crispness.
        resized = resized.filter(ImageFilter.UnsharpMask(radius=1.4, percent=90, threshold=2))
    return resized


def initials_avatar(name: str, size: int = 240) -> Image.Image:
    """Generates a clean gradient placeholder avatar from a person's initials,
    used only when no developer photo file is found."""
    initials = "".join(word[0] for word in name.split()[:2]).upper()
    img = Image.new("RGB", (size, size))
    draw = ImageDraw.Draw(img)
    top, bottom = (37, 99, 235), (29, 78, 216)
    for y in range(size):
        ratio = y / size
        color = tuple(int(top[i] + (bottom[i] - top[i]) * ratio) for i in range(3))
        draw.line([(0, y), (size, y)], fill=color)
    try:
        font = ImageFont.truetype("DejaVuSans-Bold.ttf", size=size // 2)
    except OSError:
        font = ImageFont.load_default()
    box = draw.textbbox((0, 0), initials, font=font)
    w, h = box[2] - box[0], box[3] - box[1]
    draw.text(((size - w) / 2 - box[0], (size - h) / 2 - box[1]), initials, fill="white", font=font)
    return img


def load_developer_photo() -> Image.Image:
    for path in DEVELOPER["photo_candidates"]:
        candidate = Path(path)
        if candidate.exists():
            try:
                return Image.open(candidate)
            except Exception:
                continue
    return initials_avatar(DEVELOPER["name"])


def log_event(engine: str, meta: str):
    st.session_state.history.append({
        "Time": time.strftime("%H:%M:%S"),
        "Engine": engine,
        "Details": meta,
    })


def offer_download(data: bytes, label: str, filename: str):
    st.download_button(label, data=data, file_name=filename, mime="image/png",
                        icon=":material/download:", use_container_width=True)


@contextmanager
def centered():
    """Wraps a block of Streamlit calls in the .center-col styling div."""
    st.markdown('<div class="center-col">', unsafe_allow_html=True)
    try:
        yield
    finally:
        st.markdown('</div>', unsafe_allow_html=True)


def build_text_docx(raw_text: str, content_type: str, prompt: str, language: str = None) -> bytes:
    """Renders generated Text-to-Text output into a formatted Microsoft Word
    (.docx) document -- markdown-style '#' headings become real Word heading
    styles, fenced ``` code blocks become monospace code paragraphs, and
    everything else flows as normal body paragraphs."""
    doc = Document()
    title_style = doc.styles["Title"]
    title_style.font.size = Pt(22)
    doc.add_heading("NeuralCraft AI Engine -- Generated Content", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Type: {content_type}" + (f"  |  Language: {language}" if language else "")).italic = True
    if prompt:
        p = doc.add_paragraph()
        p.add_run("Prompt: ").bold = True
        p.add_run(prompt.strip())
    doc.add_paragraph()  # spacer

    lines = raw_text.replace("\r\n", "\n").split("\n")
    in_code_block = False
    code_lines = []

    def flush_code():
        if code_lines:
            code_para = doc.add_paragraph()
            code_run = code_para.add_run("\n".join(code_lines))
            code_run.font.name = "Consolas"
            code_run.font.size = Pt(10)
            code_para.paragraph_format.left_indent = Pt(14)
            code_lines.clear()

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        if not stripped:
            continue
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            heading_text = stripped.lstrip("#").strip()
            doc.add_heading(heading_text or " ", level=max(1, level))
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            para = doc.add_paragraph(stripped)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flush_code()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def display_result(image_bytes: bytes, caption: str, filename_prefix: str, success_msg: str):
    """Shared renderer for 'generated image + success message + download button'."""
    img = Image.open(io.BytesIO(image_bytes))
    with centered():
        st.image(img, caption=caption, use_container_width=True)
    st.success(success_msg)
    st.divider()
    offer_download(image_bytes, f"Download {filename_prefix.replace('_', ' ').title()}",
                    f"{filename_prefix}_{int(time.time())}.png")


# ------------------------------------------------------------------
# SESSION STATE
# ------------------------------------------------------------------
for key, default in {
    "history": [], "image_bytes": None, "image_caption": "", "page": "home",
    "voice_transcript": "", "prompt_prefill": None,
    "text_output": "", "text_output_prompt": "", "text_output_type": "",
    "text_output_lang": None, "text_prompt_prefill": None,
    "chat_history": [],
    "video_bytes": None, "video_caption": "", "video_prompt_prefill": None,
}.items():
    st.session_state.setdefault(key, default)

# ------------------------------------------------------------------
# NAVBAR (brand + responsive navigation) + HERO
# ------------------------------------------------------------------
with st.container(key="navbar_row"):
    brand_col, nav_col, menu_col = st.columns([3, 2, 0.7], vertical_alignment="center")

    with brand_col:
        st.markdown(
            f'<div class="nav-brand"><span class="nav-logo">{brand_logo_svg("nav", 28)}</span>'
            '<span class="nav-title">NeuralCraft</span></div>',
            unsafe_allow_html=True,
        )

    with nav_col:
        # Desktop / tablet navigation -- inline links to the right of the logo.
        with st.container(key="nav_desktop"):
            nav_cols = st.columns(len(NAV_ITEMS))
            for (label, page_key, icon), col in zip(NAV_ITEMS, nav_cols):
                with col:
                    is_active = st.session_state.page == page_key
                    if st.button(label, key=f"nav_{page_key}_d", icon=icon, use_container_width=True,
                                 type="primary" if is_active else "secondary"):
                        st.session_state.page = page_key

        # Mobile navigation -- collapses into a hamburger dropdown, pinned top-right.
        # On small devices, Appearance lives inside this same menu alongside
        # Home / About / Developer instead of a separate ⋮ trigger. Recent
        # Streamlit versions keep popovers open across reruns by default, so
        # each action below explicitly closes this one via its `key` --
        # otherwise the menu would stay open after navigating or switching
        # theme instead of closing immediately as expected.
        with st.container(key="nav_mobile"):
            with st.popover("☰", key="nav_mobile_popover", use_container_width=False):
                for label, page_key, icon in NAV_ITEMS:
                    if st.button(label, key=f"nav_{page_key}_m", icon=icon, use_container_width=True):
                        st.session_state.page = page_key
                        st.session_state.nav_mobile_popover = False
                        st.rerun()
                st.markdown('<div class="mobile-appearance-block">', unsafe_allow_html=True)
                st.markdown("<div class='panel-header' style='font-size:13px;margin:10px 0 8px 0;'>Appearance</div>",
                            unsafe_allow_html=True)
                is_light_m = st.session_state.theme == "light"
                if st.button("Light Mode", key="theme_light_m", icon=":material/light_mode:", use_container_width=True,
                             type="primary" if is_light_m else "secondary"):
                    st.session_state.theme = "light"
                    st.session_state.nav_mobile_popover = False
                    st.rerun()
                if st.button("Dark Mode", key="theme_dark_m", icon=":material/dark_mode:", use_container_width=True,
                             type="primary" if not is_light_m else "secondary"):
                    st.session_state.theme = "dark"
                    st.session_state.nav_mobile_popover = False
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)

    # Appearance menu -- three-dot (⋮) popover, visible on every breakpoint,
    # holding the Light Mode / Dark Mode options (removed from the main navbar).
    # Also given a key so it closes immediately after a theme tap, same as
    # the mobile hamburger menu above.
    with menu_col, st.container(key="theme_menu_wrap"):
        with st.popover("⋮", key="theme_menu_popover", use_container_width=False, help="Appearance"):
            st.markdown("<div class='panel-header' style='font-size:14px;margin-bottom:10px;'>Appearance</div>",
                        unsafe_allow_html=True)
            is_light = st.session_state.theme == "light"
            if st.button("Light Mode", key="theme_light", icon=":material/light_mode:", use_container_width=True,
                         type="primary" if is_light else "secondary"):
                st.session_state.theme = "light"
                st.session_state.theme_menu_popover = False
                st.rerun()
            if st.button("Dark Mode", key="theme_dark", icon=":material/dark_mode:", use_container_width=True,
                         type="primary" if not is_light else "secondary"):
                st.session_state.theme = "dark"
                st.session_state.theme_menu_popover = False
                st.rerun()

st.markdown(f"""
<div class="hero-box">
    <h1 class="hero-title"><span class="typewriter">NeuralCraft Computer Vision Platform</span><span class="typewriter-cursor" aria-hidden="true"></span></h1>
    <p class="hero-subtitle">Free AI image generation, pencil-sketch conversion, voice dictation, text generation, and video generation in one lightweight app.</p>
</div>
""", unsafe_allow_html=True)

# ------------------------------------------------------------------
# VOICE-TO-PROMPT (Speech-to-Text) WIDGET
# ------------------------------------------------------------------
def render_voice_to_prompt_widget(target_label="Prompt", widget_id="default", append=False, lang="en-US", auto_detect=False):
    """Client-side mic button that uses the browser's Web Speech API to turn
    spoken words into real-time text and writes the result straight into the
    text area whose label starts with `target_label` (e.g. "Prompt" or
    "Transcript"). Runs entirely in the browser (desktop + mobile Chrome/
    Edge/Safari) -- no extra Python packages required -- and degrades
    gracefully with an inline message on unsupported browsers.

    `append=True` appends new speech onto whatever text is already in the
    field instead of overwriting it, which is what the standalone
    Speech-to-Text module uses so users can dictate in multiple takes.

    `lang` is a BCP-47 speech-recognition locale (e.g. "en-US", "ur-PK"),
    used as a sensible fallback.

    `auto_detect=True` (used by the standalone Speech-to-Text module) skips
    forcing a fixed locale altogether and instead lets the browser's speech
    engine auto-detect the spoken language from the device/browser locale --
    no language picker required, while still working the same way.
    """
    btn_id = f"nc-mic-btn-{widget_id}"
    status_id = f"nc-mic-status-{widget_id}"
    mic_html = f"""
    <style>
        body {{ margin: 0; }}
        .mic-btn-wrap {{ display: flex; align-items: center; gap: 8px; margin: 4px 0 2px 0;
                        font-family: 'Plus Jakarta Sans', -apple-system, sans-serif; }}
        .mic-btn {{
            display: inline-flex; align-items: center; justify-content: center;
            width: 40px; height: 40px; min-width: 40px; border-radius: 50%; padding: 0;
            background: linear-gradient(135deg, #2563eb, #1d4ed8);
            border: none; cursor: pointer;
            box-shadow: 0 4px 14px rgba(37,99,235,.35);
            transition: transform .2s cubic-bezier(.4,0,.2,1), box-shadow .2s cubic-bezier(.4,0,.2,1), filter .2s cubic-bezier(.4,0,.2,1);
        }}
        .mic-btn svg {{ width: 19px; height: 19px; stroke: #fff; fill: none; pointer-events: none; }}
        .mic-btn:hover {{ transform: scale(1.08); box-shadow: 0 8px 20px rgba(37,99,235,.45); filter: brightness(1.08); }}
        .mic-btn:active {{ transform: scale(.95); }}
        .mic-btn.listening {{ animation: mic-pulse 1.1s ease-in-out infinite; background: linear-gradient(135deg, #ef4444, #dc2626); }}
        @keyframes mic-pulse {{
            0%   {{ box-shadow: 0 0 0 0 rgba(239,68,68,.55); }}
            70%  {{ box-shadow: 0 0 0 12px rgba(239,68,68,0); }}
            100% {{ box-shadow: 0 0 0 0 rgba(239,68,68,0); }}
        }}
        .mic-status {{ font-size: 12.5px; color: #a3adc2; line-height: 1.4; }}
    </style>
    <div class="mic-btn-wrap">
      <button id="{btn_id}" class="mic-btn" type="button" title="Speak your prompt" aria-label="Speak your prompt">{ICONS['mic']}</button>
      <span id="{status_id}" class="mic-status">Click the mic and speak -- your words appear as editable text in real time</span>
    </div>
    <script>
    (function() {{
        const btn = document.getElementById('{btn_id}');
        const status = document.getElementById('{status_id}');
        const targetLabel = {target_label!r};
        const appendMode = {str(append).lower()};
        const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
        if (!SpeechRecognition) {{
            status.textContent = "Voice input isn't supported in this browser. Try Chrome, Edge, or Safari.";
            btn.disabled = true;
            btn.style.opacity = '0.5';
            btn.style.cursor = 'not-allowed';
            return;
        }}
        const recognition = new SpeechRecognition();
        const autoDetect = {str(auto_detect).lower()};
        if (!autoDetect) {{
            recognition.lang = {lang!r};
        }} else if (navigator.language) {{
            // Auto mode: seed with the browser/device locale (best available
            // signal for language auto-detection in the Web Speech API),
            // rather than a fixed picker value.
            recognition.lang = navigator.language;
        }}
        recognition.interimResults = true;
        recognition.continuous = false;
        let listening = false;
        let baseText = '';

        function findTarget() {{
            const doc = window.parent.document;
            const areas = doc.querySelectorAll('[data-testid="stTextArea"] textarea');
            let target = null;
            areas.forEach(function(el) {{
                const wrap = el.closest('[data-testid="stTextArea"]');
                const label = wrap ? wrap.querySelector('label') : null;
                if (label && label.innerText.trim().indexOf(targetLabel) === 0) target = el;
            }});
            if (!target && areas.length) target = areas[0];
            return target;
        }}

        function setFieldValue(text) {{
            try {{
                const target = findTarget();
                if (!target) return false;
                const setter = Object.getOwnPropertyDescriptor(window.parent.HTMLTextAreaElement.prototype, 'value').set;
                setter.call(target, text);
                target.dispatchEvent(new Event('input', {{ bubbles: true }}));
                target.focus();
                target.blur();
                return true;
            }} catch (err) {{
                return false;
            }}
        }}

        btn.addEventListener('click', function() {{
            if (listening) {{ recognition.stop(); return; }}
            if (appendMode) {{
                const t = findTarget();
                baseText = t && t.value ? (t.value.trim() + ' ') : '';
            }}
            try {{ recognition.start(); }} catch (err) {{ /* already running */ }}
        }});
        recognition.onstart = function() {{
            listening = true;
            btn.classList.add('listening');
            status.textContent = 'Listening... speak now';
        }};
        recognition.onresult = function(event) {{
            let transcript = '';
            for (let i = 0; i < event.results.length; i++) {{
                transcript += event.results[i][0].transcript;
            }}
            status.textContent = baseText + transcript;
            if (event.results[event.results.length - 1].isFinal) {{
                const clean = (baseText + transcript).trim();
                const ok = setFieldValue(clean);
                status.textContent = ok ? ('Updated: "' + clean + '"')
                                         : 'Could not reach the text field -- please paste manually.';
            }}
        }};
        recognition.onerror = function(event) {{
            status.textContent = 'Voice input error: ' + event.error + '. Please try again.';
        }};
        recognition.onend = function() {{
            listening = false;
            btn.classList.remove('listening');
        }};
    }})();
    </script>
    """
    components.html(mic_html, height=50)


# ------------------------------------------------------------------
# HOME PAGE HELPERS -- TEXT-TO-IMAGE (prompt only)
# ------------------------------------------------------------------
def render_text_to_image(col_control, col_canvas):
    with col_control, st.container(border=True):
        icon_header("settings", "Model Control Dashboard")
        st.info("Describe your idea, tune the style and resolution, then generate.")

        preset = st.selectbox("Quick Idea Starter", list(PROMPT_PRESETS.keys()))
        render_voice_to_prompt_widget(target_label="Prompt", widget_id="prompt")
        prefill = st.session_state.pop("prompt_prefill", None)
        prompt = st.text_area(
            "Prompt", value=prefill if prefill else PROMPT_PRESETS[preset],
            placeholder="e.g., A cute fluffy white Persian cat playing with a red wool ball...",
            height=100,
        )
        st.divider()
        style = st.selectbox("Style Preset", STYLE_PRESETS)
        ratio_label = st.selectbox("Resolution", list(ASPECT_RATIOS.keys()))
        width, height = ASPECT_RATIOS[ratio_label]

        st.divider()
        c1, c2 = st.columns([2, 1])
        generate = c1.button("Generate Image", icon=":material/bolt:", use_container_width=True, type="primary")
        if c2.button("Clear", icon=":material/delete:", use_container_width=True):
            st.session_state.image_bytes = None
            st.session_state.image_caption = ""
            st.rerun()

        if generate and prompt.strip():
            status_ph = st.empty()

            def _report_progress(attempt, max_attempts, message):
                status_ph.caption(f"⏳ Attempt {attempt}/{max_attempts} -- {message}")

            with st.spinner("Generating your image... this can take a few seconds"):
                image_bytes, error = generate_ai_image(
                    prompt, width, height, style, progress_cb=_report_progress,
                )
            status_ph.empty()
            if error:
                st.error(error)
            else:
                st.session_state.image_bytes = image_bytes
                st.session_state.image_caption = prompt
                log_event("Flux (Pollinations.ai)", f"{width}x{height} | {style}")

    with col_canvas, st.container(border=True):
        icon_header("image", "Output Canvas")
        st.caption("Your generated image will appear here.")
        st.divider()

        if st.session_state.image_bytes:
            display_result(st.session_state.image_bytes, st.session_state.image_caption,
                            "neuralcraft", "🎉 Image generated successfully!")
        else:
            st.warning("Provide a prompt to see results here.")


# ------------------------------------------------------------------
# HOME PAGE HELPERS -- PHOTO-TO-SKETCH
# ------------------------------------------------------------------
def render_sketch_studio(col_control, col_canvas):
    with col_control, st.container(border=True):
        icon_header("pencil", "Sketch Configurator")
        st.info("Choose an input source below. Shading controls sit right beside the "
                 "sketch preview so you can fine-tune it in real time.")

        input_mode = st.radio(
            "Input Source",
            [":material/upload_file: Upload File", ":material/photo_camera: Use Camera"],
            horizontal=True,
        )
        st.divider()

        source_image, source_name = None, ""
        if input_mode == ":material/upload_file: Upload File":
            uploaded = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png"])
            if uploaded is not None:
                source_image = Image.open(uploaded)
                source_name = uploaded.name
        else:
            camera_file = st.camera_input("Take a photo")
            if camera_file is not None:
                source_image = Image.open(camera_file)
                source_name = "Camera Snapshot"

    with col_canvas, st.container(border=True):
        icon_header("image", "Sketch Canvas")
        st.caption("Original image, live shading controls, and the converted sketch appear below.")
        st.divider()

        if source_image is not None:
            st.markdown('<div class="preview-box">', unsafe_allow_html=True)
            with centered():
                st.image(source_image, caption="Original Image", use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
            st.divider()

            # Shading Control column sits immediately to the LEFT of the sketch
            # preview column, so adjustments are visible against the live result.
            shading_col, sketch_col = st.columns([1, 2.2], gap="medium")
            with shading_col:
                st.markdown('<div class="shading-panel">', unsafe_allow_html=True)
                st.markdown('<div class="density-label">🎚️ Density</div>', unsafe_allow_html=True)
                blur = st.slider("Sketch Density", 3, 99, 25, step=2, label_visibility="collapsed",
                                  help="Lower = sharper lines. Higher = smoother, richer shading.")
                st.caption("Lower = sharper lines. Higher = smoother shading.")

                st.markdown('<div class="shading-label">🌓 Shading Depth</div>', unsafe_allow_html=True)
                shade = st.slider("Shading Depth", 0.0, 1.0, 0.5, step=0.1, label_visibility="collapsed")

                st.markdown('<div class="shading-label">✏️ Stroke Contrast</div>', unsafe_allow_html=True)
                contrast = st.slider("Stroke Contrast", 1.0, 2.0, 1.1, step=0.1, label_visibility="collapsed")
                st.markdown('</div>', unsafe_allow_html=True)

            with st.spinner("Converting to pencil sketch..."):
                sketch = to_pencil_sketch(source_image, blur, shade, contrast)
            buf = io.BytesIO()
            sketch.save(buf, format="PNG")
            sketch_bytes = buf.getvalue()

            with sketch_col:
                st.markdown('<div class="preview-box preview-box--sketch">', unsafe_allow_html=True)
                with centered():
                    st.image(sketch, caption="Pencil Sketch Result", use_container_width=True)
                st.success("📊 Sketch conversion complete!")
                st.divider()

                st.markdown('<div class="density-label">📥 Download Sketch</div>', unsafe_allow_html=True)
                size_cols = st.columns(len(DOWNLOAD_SIZES))
                for (size_label, max_w), size_col in zip(DOWNLOAD_SIZES.items(), size_cols):
                    with size_col:
                        sized = resize_for_download(sketch, max_w)
                        sized_buf = io.BytesIO()
                        sized.save(sized_buf, format="PNG")
                        st.download_button(
                            f"{size_label}", data=sized_buf.getvalue(),
                            file_name=f"sketch_{size_label.lower()}_{int(time.time())}.png",
                            mime="image/png", icon=":material/download:", use_container_width=True,
                            key=f"sketch_dl_{size_label.lower()}",
                        )
                st.markdown('</div>', unsafe_allow_html=True)

            log_event(f"OpenCV Sketch ({input_mode})", f"Source: {source_name} | Density: {blur}")
        else:
            st.warning("Upload a file or take a photo to begin.")


# ------------------------------------------------------------------
# HOME PAGE HELPERS -- SPEECH-TO-TEXT
# ------------------------------------------------------------------
def render_speech_to_text(col_control, col_canvas):
    with col_control, st.container(border=True):
        icon_header("mic", "Voice Capture")
        st.info("Click the microphone and speak. Your words are converted to editable "
                 "text in real time -- pause and click again to keep adding to it.")
        render_voice_to_prompt_widget(target_label="Transcript", widget_id="stt", append=True,
                                       auto_detect=True)
        st.caption("Runs entirely in your browser (Chrome, Edge, or Safari work best) and "
                    "automatically detects your spoken language from your device settings. "
                    "Grant microphone permission when prompted -- nothing is uploaded.")
        st.divider()
        if st.button("Clear Transcript", icon=":material/delete:", use_container_width=True):
            st.session_state.voice_transcript = ""
            st.rerun()

    with col_canvas, st.container(border=True):
        icon_header("text", "Transcript")
        st.caption("Edit freely, copy it out, or send it straight into the Text-to-Image prompt.")
        st.divider()

        transcript = st.text_area(
            "Transcript", value=st.session_state.voice_transcript, height=180,
            placeholder="Your speech will appear here as editable text...",
            label_visibility="collapsed",
        )
        st.session_state.voice_transcript = transcript

        if transcript.strip():
            st.divider()
            st.markdown('<div class="density-label">📥 Copy Text</div>', unsafe_allow_html=True)
            st.code(transcript, language=None)

            c1, c2 = st.columns(2)
            with c1:
                if st.button("Use as Prompt", icon=":material/auto_awesome:", use_container_width=True, type="primary"):
                    st.session_state.prompt_prefill = transcript.strip()
                    st.session_state.mode = MODE_TEXT_TO_IMAGE
                    log_event("Speech-to-Text", f"Sent {len(transcript.strip())} chars to prompt")
                    st.rerun()
            with c2:
                if st.download_button(
                    "Download .txt", data=transcript.strip(), file_name=f"transcript_{int(time.time())}.txt",
                    mime="text/plain", icon=":material/download:", use_container_width=True,
                ):
                    log_event("Speech-to-Text", f"Downloaded {len(transcript.strip())} chars")
        else:
            st.warning("Use the microphone on the left to start dictating.")


# ------------------------------------------------------------------
# HOME PAGE HELPERS -- TEXT-TO-TEXT
# ------------------------------------------------------------------
def render_text_to_text(col_control, col_canvas):
    with col_control, st.container(border=True):
        icon_header("settings", "Content Control Dashboard")
        st.info("Describe what you need -- theory content like an assignment, report, "
                 "or explanation, or programming code -- then generate.")

        content_type = st.selectbox("Content Type", ["Theory", "Programming Code"])
        language = None
        if content_type == "Programming Code":
            language = st.selectbox("Programming Language", CODE_LANGUAGES)

        preset = st.selectbox("Quick Idea Starter", list(TEXT_PRESETS.keys()))
        render_voice_to_prompt_widget(target_label="Text Prompt", widget_id="text2text")
        prefill = st.session_state.pop("text_prompt_prefill", None)
        prompt = st.text_area(
            "Text Prompt", value=prefill if prefill else TEXT_PRESETS[preset],
            placeholder="e.g., Write a report on renewable energy trends, or a Python "
                        "function that sorts a list of dictionaries by a key...",
            height=110,
        )
        st.divider()
        c1, c2 = st.columns([2, 1])
        generate = c1.button("Generate Text", icon=":material/bolt:", use_container_width=True, type="primary")
        if c2.button("Clear", icon=":material/delete:", use_container_width=True, key="clear_text2text"):
            st.session_state.text_output = ""
            st.session_state.text_output_prompt = ""
            st.rerun()

        # A failed attempt sets this flag so a single tap on "Retry Generation"
        # re-runs the exact same request without the user needing to press
        # "Generate Text" again or retype anything.
        do_generate = generate or st.session_state.pop("text_retry_pending", False)

        if do_generate and prompt.strip():
            with st.spinner("Generating your content... long-form content can take up to a minute or two"):
                text_out, error = generate_ai_text(prompt, content_type, language)

            if error:
                st.error(f"⚠️ {error}")
                if st.button("🔁 Retry Generation", key="retry_text2text", use_container_width=True):
                    st.session_state.text_retry_pending = True
                    st.rerun()
            else:
                st.session_state.text_output = text_out
                st.session_state.text_output_prompt = prompt
                st.session_state.text_output_type = content_type
                st.session_state.text_output_lang = language
                log_event("Text-to-Text (Pollinations)",
                          f"{content_type}" + (f" | {language}" if language else ""))

    with col_canvas, st.container(border=True):
        icon_header("text", "Generated Output")
        st.caption("Your AI-generated text will appear here.")
        st.divider()

        if st.session_state.text_output:
            out_type = st.session_state.text_output_type
            out_lang = st.session_state.text_output_lang
            raw_text = st.session_state.text_output

            if out_type == "Programming Code":
                code_only = _strip_code_fences(raw_text)
                st.code(code_only, language=CODE_HIGHLIGHT_LANG.get(out_lang, None))
            else:
                with centered():
                    st.markdown(f'<div class="preview-box" style="text-align:left;padding:20px 22px;">',
                                unsafe_allow_html=True)
                    st.markdown(raw_text)
                    st.markdown('</div>', unsafe_allow_html=True)

            st.success("🎉 Content generated successfully!")
            st.divider()

            c1, c2 = st.columns(2)
            if out_type == "Programming Code":
                # Download the raw code with the extension matching the
                # selected language (.py, .cpp, .java, .js, .sql, etc.)
                # instead of wrapping it in a generic .docx file.
                code_ext = CODE_FILE_EXT.get(out_lang, "txt")
                with c1:
                    st.download_button(
                        f"Download .{code_ext}", data=code_only,
                        file_name=f"neuralcraft_code_{int(time.time())}.{code_ext}",
                        mime="text/plain",
                        icon=":material/download:", use_container_width=True,
                    )
            else:
                docx_bytes = build_text_docx(
                    raw_text, out_type, st.session_state.text_output_prompt, out_lang,
                )
                with c1:
                    st.download_button(
                        "Download .docx", data=docx_bytes,
                        file_name=f"neuralcraft_text_{int(time.time())}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        icon=":material/download:", use_container_width=True,
                    )
            with c2:
                if st.button("Use as Image Prompt", icon=":material/auto_awesome:", use_container_width=True):
                    st.session_state.prompt_prefill = raw_text.strip()[:400]
                    st.session_state.mode = MODE_TEXT_TO_IMAGE
                    log_event("Text-to-Text", f"Sent {len(raw_text.strip())} chars to image prompt")
                    st.rerun()

            with st.expander("📋 Copy Raw Text"):
                st.code(raw_text, language=None)
        else:
            st.warning("Provide a prompt to see results here.")


# ------------------------------------------------------------------
# HOME PAGE HELPERS -- TEXT-TO-VIDEO
# ------------------------------------------------------------------
def render_text_to_video(col_control, col_canvas):
    with col_control, st.container(border=True):
        icon_header("settings", "Video Generation Console")
        st.info("Describe a scene or action, tune duration and orientation, then "
                 "generate a short AI video (up to 10 seconds, model permitting).")

        preset = st.selectbox("Quick Idea Starter", list(VIDEO_PROMPT_PRESETS.keys()), key="video_preset")
        render_voice_to_prompt_widget(target_label="Video Prompt", widget_id="video")
        prefill = st.session_state.pop("video_prompt_prefill", None)
        prompt = st.text_area(
            "Video Prompt", value=prefill if prefill else VIDEO_PROMPT_PRESETS[preset],
            placeholder="e.g., A paper airplane gliding through a sunlit office, slow motion...",
            height=100,
        )
        st.divider()
        model_label = st.selectbox("Video Model", list(VIDEO_MODELS.keys()))
        d1, d2 = st.columns(2)
        min_dur, max_dur = VIDEO_DURATION_RANGE
        duration = d1.slider(
            "Duration (seconds)", min_dur, max_dur, 5,
            help="Actual length may be adjusted to the underlying model's supported range.",
        )
        aspect_label = d2.selectbox("Orientation", list(VIDEO_ASPECTS.keys()))
        aspect_ratio = VIDEO_ASPECTS[aspect_label]

        st.divider()
        c1, c2 = st.columns([2, 1])
        generate = c1.button("Generate Video", icon=":material/movie:", use_container_width=True,
                              type="primary", key="gen_video_btn")
        if c2.button("Clear", icon=":material/delete:", use_container_width=True, key="clear_video"):
            st.session_state.video_bytes = None
            st.session_state.video_caption = ""
            st.rerun()

        # A failed attempt sets this flag so a single tap on "Retry
        # Generation" re-runs the exact same request without retyping.
        do_generate = generate or st.session_state.pop("video_retry_pending", False)

        if do_generate and prompt.strip():
            progress_bar = st.progress(0, text="Preparing request...")

            def _report_progress(fraction, message):
                progress_bar.progress(min(1.0, max(0.0, fraction)), text=message)

            video_bytes, error = generate_ai_video(
                prompt, duration, aspect_ratio, model_label, progress_cb=_report_progress,
            )
            progress_bar.empty()

            if error:
                st.error(f"⚠️ {error}")
                if st.button("🔁 Retry Generation", key="retry_video", use_container_width=True):
                    st.session_state.video_retry_pending = True
                    st.rerun()
            else:
                st.session_state.video_bytes = video_bytes
                st.session_state.video_caption = prompt
                log_event("Text-to-Video (Pollinations)",
                          f"{model_label} | {duration}s | {aspect_label}")

    with col_canvas, st.container(border=True):
        icon_header("video", "Video Preview")
        st.caption("Your generated video will appear here.")
        st.divider()

        if st.session_state.video_bytes:
            with centered():
                st.video(st.session_state.video_bytes)
            st.success("🎬 Video generated successfully!")
            st.divider()
            st.download_button(
                "Download Video", data=st.session_state.video_bytes,
                file_name=f"neuralcraft_video_{int(time.time())}.mp4",
                mime="video/mp4", icon=":material/download:", use_container_width=True,
                key="video_dl",
            )
        else:
            st.warning("Provide a prompt to see results here.")


# ------------------------------------------------------------------
# HOME PAGE HELPERS -- AI CHATBOT
# ------------------------------------------------------------------
def render_ai_chatbot():
    """Standalone conversational AI assistant -- its own top-level module
    (previously embedded inside Text-to-Text). Rendered full-width since a
    chat thread doesn't naturally split into a control/canvas pair like the
    other five modules."""
    with st.container(border=True):
        icon_header("chat", "AI Chatbot")
        st.caption("Ask the AI assistant anything -- independent of the other AI modules.")
        st.divider()

        for msg in st.session_state.chat_history:
            avatar = ":material/smart_toy:" if msg["role"] == "assistant" else ":material/person:"
            with st.chat_message(msg["role"], avatar=avatar):
                st.markdown(msg["content"])

        if not st.session_state.chat_history:
            st.info("Start a conversation below -- ask a question and get an instant AI-generated response.")
        elif st.button("Clear Chat", icon=":material/delete:", key="clear_chat", use_container_width=True):
            st.session_state.chat_history = []
            st.rerun()

        user_msg = st.chat_input("Ask the AI assistant anything...")
        if user_msg:
            st.session_state.chat_history.append({"role": "user", "content": user_msg})
            with st.chat_message("user", avatar=":material/person:"):
                st.markdown(user_msg)
            with st.chat_message("assistant", avatar=":material/smart_toy:"):
                with st.spinner("Thinking..."):
                    # Only the most recent turns are sent as context -- plenty
                    # for a coherent conversation without an ever-growing payload.
                    reply, error = generate_chat_response(st.session_state.chat_history[-12:])
                if error:
                    st.error(f"⚠️ {error}")
                    if st.button("Retry", icon=":material/refresh:", key="retry_chat", use_container_width=True):
                        st.session_state.chat_retry_pending = True
                        st.rerun()
                else:
                    st.markdown(reply)
                    st.session_state.chat_history.append({"role": "assistant", "content": reply})
        elif st.session_state.pop("chat_retry_pending", False) and st.session_state.chat_history \
                and st.session_state.chat_history[-1]["role"] == "user":
            with st.chat_message("assistant", avatar=":material/smart_toy:"):
                with st.spinner("Thinking..."):
                    reply, error = generate_chat_response(st.session_state.chat_history[-12:])
                if error:
                    st.error(f"⚠️ {error}")
                else:
                    st.markdown(reply)
                    st.session_state.chat_history.append({"role": "assistant", "content": reply})


# ------------------------------------------------------------------
# HOME PAGE
# ------------------------------------------------------------------
def render_home():
    st.markdown('<div class="mode-switch-wrap">', unsafe_allow_html=True)
    mode = st.radio(
        "Mode",
        [MODE_TEXT_TO_IMAGE, MODE_PHOTO_TO_SKETCH, MODE_SPEECH_TO_TEXT, MODE_TEXT_TO_TEXT,
         MODE_TEXT_TO_VIDEO, MODE_AI_CHATBOT],
        horizontal=True, label_visibility="collapsed", key="mode",
    )
    st.markdown('</div>', unsafe_allow_html=True)

    with st.container(key="studio_columns"):
        if mode == MODE_AI_CHATBOT:
            # The chatbot is a single conversational thread, not a
            # control/canvas pair, so it renders full-width on its own.
            render_ai_chatbot()
        else:
            col_control, col_canvas = st.columns([1, 1.2], gap="large")
            if mode == MODE_TEXT_TO_IMAGE:
                render_text_to_image(col_control, col_canvas)
            elif mode == MODE_PHOTO_TO_SKETCH:
                render_sketch_studio(col_control, col_canvas)
            elif mode == MODE_SPEECH_TO_TEXT:
                render_speech_to_text(col_control, col_canvas)
            elif mode == MODE_TEXT_TO_TEXT:
                render_text_to_text(col_control, col_canvas)
            else:
                render_text_to_video(col_control, col_canvas)

    st.divider()
    with st.expander("Session Activity Log", icon=":material/monitoring:", expanded=True):
        st.caption("Log of generation and conversion events from this session.")
        if st.session_state.history:
            st.dataframe(pd.DataFrame(st.session_state.history), use_container_width=True)
        else:
            st.caption("No activity yet. Generate an image or convert a sketch to populate this log.")


# ------------------------------------------------------------------
# ABOUT PAGE
# ------------------------------------------------------------------
def render_about():
    icon_header("info", "About NeuralCraft")
    st.write(
        "NeuralCraft is a lightweight, AI-powered creative studio built entirely on "
        "free and open-source models. It turns your words -- and your photos -- into "
        "new images, sketches, and written content in seconds."
    )
    st.markdown(f"""
    <div class="feature-grid--primary">
        <div class="feature-card">
            <div class="feature-icon">{ICONS['image']}</div>
            <div class="feature-title">Text-to-Image</div>
            <div class="feature-text">Describe any scene, style, or idea in plain
            language and NeuralCraft generates a brand-new image from scratch using
            a free open text-to-image model -- no design skills required.</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">{ICONS['pencil']}</div>
            <div class="feature-title">Photo-to-Sketch</div>
            <div class="feature-text">Convert any photo into a sharp, professional
            pencil-sketch rendering with clean hand-drawn linework, natural shading,
            and adjustable detail -- all processed locally with OpenCV.</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">{ICONS['mic']}</div>
            <div class="feature-title">Speech-to-Text</div>
            <div class="feature-text">Dictate instead of typing -- click the mic and
            watch your words appear as editable text in real time, ready to copy,
            refine, or send straight into the Text-to-Image prompt.</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">{ICONS['text']}</div>
            <div class="feature-title">Text-to-Text</div>
            <div class="feature-text">Generate theory content like assignments,
            reports, and explanations, or ready-to-run code in Python, C++, Java,
            JavaScript, SQL and more -- just describe what you need.</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">{ICONS['video']}</div>
            <div class="feature-title">Text-to-Video</div>
            <div class="feature-text">Describe a scene or action and NeuralCraft
            generates a short AI video clip up to 10 seconds long -- choose the
            orientation and duration, preview it instantly, and download the
            finished MP4.</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">{ICONS['chat']}</div>
            <div class="feature-title">AI Chatbot</div>
            <div class="feature-text">Chat with a friendly, knowledgeable AI
            assistant for quick answers and conversation -- a standalone
            module, always available independent of the other five AI tools.</div>
        </div>
    </div>
    <div class="feature-grid">
        <div class="feature-card">
            <div class="feature-icon">{ICONS['bolt']}</div>
            <div class="feature-title">Fast &amp; Lightweight</div>
            <div class="feature-text">No heavyweight local models or GPU required --
            NeuralCraft runs entirely on free hosted inference and efficient
            on-device image processing.</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ------------------------------------------------------------------
# DEVELOPER PAGE
# ------------------------------------------------------------------
def render_developer():
    icon_header("developer", "Meet the Developer")
    photo = load_developer_photo()
    buf = io.BytesIO()
    photo.convert("RGB").save(buf, format="PNG")
    photo_b64 = base64.b64encode(buf.getvalue()).decode()

    st.markdown(f"""
    <div class="profile-wrap">
        <div class="profile-card">
            <div class="profile-photo-wrap">
                <img class="profile-photo" src="data:image/png;base64,{photo_b64}" />
            </div>
            <div class="profile-info">
                <p class="profile-name">{DEVELOPER['name']}</p>
                <p class="profile-role">{DEVELOPER['profession']}</p>
                <p class="profile-detail"><span class="profile-detail-icon">{ICONS['graduation']}</span>{DEVELOPER['degree']}</p>
                <p class="profile-detail"><span class="profile-detail-icon">{ICONS['idea']}</span>Creator of the NeuralCraft AI Engine</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if not any(Path(p).exists() for p in DEVELOPER["photo_candidates"]):
        st.caption(
            "ℹ️ No photo file found -- showing a placeholder avatar. "
            "Add your photo as `assets/developer.jpg` next to this script to display it."
        )


# ------------------------------------------------------------------
# PAGE ROUTER
# ------------------------------------------------------------------
PAGES = {"home": render_home, "about": render_about, "dev": render_developer}
PAGES.get(st.session_state.page, render_home)()
